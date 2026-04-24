from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ───────────────────────────────────────────────────────────
SWIFT_BLUE   = RGBColor(0x00, 0x33, 0x72)
SWIFT_LIGHT  = RGBColor(0x00, 0x70, 0xC0)
SWIFT_GOLD   = RGBColor(0xF0, 0xA0, 0x00)
SWIFT_GREEN  = RGBColor(0x00, 0x70, 0x50)
SWIFT_RED    = RGBColor(0xC0, 0x00, 0x00)
SWIFT_TEAL   = RGBColor(0x00, 0x8B, 0x8B)
SWIFT_CLOUD  = RGBColor(0x00, 0x6A, 0xB0)
PURPLE       = RGBColor(0x70, 0x30, 0xA0)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY   = RGBColor(0xF2, 0xF2, 0xF2)
MID_GREY     = RGBColor(0xBF, 0xBF, 0xBF)
DARK_GREY    = RGBColor(0x26, 0x26, 0x26)
CODE_BG      = RGBColor(0x1E, 0x1E, 0x2E)
CODE_FG      = RGBColor(0xA6, 0xE3, 0xA1)

# ── Helpers ──────────────────────────────────────────────────────────────────
def rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex(rgb))
    tcPr.append(shd)

def set_no_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   'none')
        tag.set(qn('w:sz'),    '0')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), 'auto')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_light_border(cell, color='BFBFBF'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def heading(doc, text, color=SWIFT_BLUE, size=13, bold=True, space_before=14, space_after=4):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.name      = 'Calibri'
    return p

def body(doc, text, color=DARK_GREY, size=9.5, italic=False, bold=False, space_before=2, space_after=2):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size      = Pt(size)
    run.font.italic    = italic
    run.font.bold      = bold
    run.font.name      = 'Calibri'
    return p

def code_block(doc, lines):
    """Dark-background monospace code block."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, CODE_BG)
    set_no_border(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Pt(8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run('\n')
        run = p.add_run(line)
        run.font.name      = 'Courier New'
        run.font.size      = Pt(8.5)
        run.font.color.rgb = CODE_FG
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def divider(doc, color='003372', sz='8'):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def banner(doc, text, bg=SWIFT_BLUE, fg=WHITE, size=12.5):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    set_no_border(cell)
    cell.width = Inches(6.47)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(7)
    p.paragraph_format.left_indent  = Pt(10)
    run = p.add_run(text)
    run.font.color.rgb = fg
    run.font.size      = Pt(size)
    run.font.bold      = True
    run.font.name      = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def step_box(doc, step_num, title, description, color=SWIFT_BLUE):
    tbl  = doc.add_table(rows=1, cols=2)
    tbl.alignment    = WD_TABLE_ALIGNMENT.LEFT
    tbl.allow_autofit = False
    badge = tbl.cell(0, 0)
    cont  = tbl.cell(0, 1)
    badge.width = Inches(0.62)
    cont.width  = Inches(5.85)
    set_cell_bg(badge, color)
    set_cell_bg(cont,  LIGHT_GREY)
    set_light_border(badge, rgb_hex(color))
    set_light_border(cont,  'D0D0D0')
    bp = badge.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(8)
    br = bp.add_run(f"STEP\n{step_num}")
    br.font.color.rgb = WHITE
    br.font.size      = Pt(8.5)
    br.font.bold      = True
    br.font.name      = 'Calibri'
    cp = cont.paragraphs[0]
    cp.paragraph_format.left_indent  = Pt(8)
    cp.paragraph_format.space_before = Pt(5)
    cp.paragraph_format.space_after  = Pt(5)
    ct = cp.add_run(f"{title}\n")
    ct.font.color.rgb = color
    ct.font.size      = Pt(10.5)
    ct.font.bold      = True
    ct.font.name      = 'Calibri'
    cd = cp.add_run(description)
    cd.font.color.rgb = DARK_GREY
    cd.font.size      = Pt(9)
    cd.font.name      = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def note_box(doc, label, text, label_color=SWIFT_GOLD):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = tbl.cell(0, 0)
    rc = tbl.cell(0, 1)
    lc.width = Inches(1.1)
    rc.width = Inches(5.37)
    set_cell_bg(lc, label_color)
    set_cell_bg(rc, LIGHT_GREY)
    set_no_border(lc)
    set_no_border(rc)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(5)
    lr = lp.add_run(label)
    lr.font.color.rgb = WHITE
    lr.font.bold      = True
    lr.font.size      = Pt(8.5)
    lr.font.name      = 'Calibri'
    rp = rc.paragraphs[0]
    rp.paragraph_format.left_indent  = Pt(6)
    rp.paragraph_format.space_before = Pt(5)
    rp.paragraph_format.space_after  = Pt(5)
    rr = rp.add_run(text)
    rr.font.size = Pt(9)
    rr.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def grid_table(doc, headers, rows, col_widths=None, hdr_color=SWIFT_BLUE):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, hdr_color)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.font.color.rgb = WHITE
        run.font.bold      = True
        run.font.size      = Pt(9)
        run.font.name      = 'Calibri'
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = WHITE if ri % 2 == 0 else LIGHT_GREY
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Pt(3)
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = 'Calibri'
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def flow_row(doc, phases):
    """Horizontal flow diagram as a table row."""
    n   = len(phases)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, color) in enumerate(phases):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, color)
        set_no_border(cell)
        cell.width = Inches(6.47 / n)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(8)
        run = p.add_run(label)
        run.font.color.rgb = WHITE
        run.font.size      = Pt(7.5)
        run.font.bold      = True
        run.font.name      = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def arch_diagram(doc, rows_data, title=None):
    """Simple visual architecture diagram as a table grid."""
    if title:
        heading(doc, title, SWIFT_BLUE, 10, space_before=6)
    tbl = doc.add_table(rows=len(rows_data), cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (text, bg, fg) in enumerate(rows_data):
        cell = tbl.cell(i, 0)
        set_cell_bg(cell, bg)
        set_no_border(cell)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        run = p.add_run(text)
        run.font.color.rgb = fg
        run.font.size      = Pt(9)
        run.font.bold      = True
        run.font.name      = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT BUILD
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.27)
sec.page_height   = Inches(11.69)
sec.left_margin   = Inches(0.85)
sec.right_margin  = Inches(0.85)
sec.top_margin    = Inches(0.65)
sec.bottom_margin = Inches(0.65)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ─── COVER PAGE ──────────────────────────────────────────────────────────────
cover_tbl = doc.add_table(rows=1, cols=1)
cc = cover_tbl.cell(0, 0)
set_cell_bg(cc, SWIFT_BLUE)
set_no_border(cc)
cc.width = Inches(6.47)
for txt, fg, sz, bold, sb, sa in [
    ("SWIFT BIC ONBOARDING GUIDE",                WHITE,      22, True,  36, 4),
    ("Cloud Edition — Including API Integration", SWIFT_GOLD, 13, False, 2,  4),
    ("v2.0 | April 2026 | CONFIDENTIAL",         MID_GREY,   9,  False, 2,  32),
]:
    p = cc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(txt)
    r.font.color.rgb = fg
    r.font.size      = Pt(sz)
    r.font.bold      = bold
    r.font.name      = 'Calibri'
doc.add_paragraph()

meta_tbl = doc.add_table(rows=2, cols=4)
meta_data = [
    ("Document Owner",  "SWIFT Security Officer"),
    ("Classification",  "CONFIDENTIAL"),
    ("Version",         "2.0"),
    ("Effective Date",  "April 2026"),
    ("Scope",           "Cross-Border Transactions + Cloud + API"),
    ("Compliance",      "SWIFT CSP / ISO 20022 / CBPR+"),
    ("Review Cycle",    "Annual"),
    ("Status",          "APPROVED"),
]
for idx, (label, value) in enumerate(meta_data):
    ri, ci = divmod(idx, 4)
    cell   = meta_tbl.cell(ri, ci)
    set_cell_bg(cell, SWIFT_BLUE if ri == 0 else LIGHT_GREY)
    set_no_border(cell)
    p   = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(4)
    rl = p.add_run(f"{label}\n")
    rl.font.size      = Pt(7)
    rl.font.bold      = True
    rl.font.color.rgb = WHITE if ri == 0 else SWIFT_BLUE
    rl.font.name      = 'Calibri'
    rv = p.add_run(value)
    rv.font.size      = Pt(8.5)
    rv.font.bold      = False
    rv.font.color.rgb = SWIFT_GOLD if ri == 0 else DARK_GREY
    rv.font.name      = 'Calibri'
doc.add_paragraph()

# ─── SECTION 01 – OVERVIEW ───────────────────────────────────────────────────
banner(doc, "01  |  OVERVIEW & SCOPE")
heading(doc, "What is a BIC (SWIFT Code)?", SWIFT_BLUE, 11, space_before=4)
body(doc, (
    "A Bank Identifier Code (BIC) — also called a SWIFT Code — is an 8 or 11-character "
    "alphanumeric identifier assigned by SWIFT (Society for Worldwide Interbank Financial "
    "Telecommunication) to uniquely identify a financial institution in cross-border transactions."
))
doc.add_paragraph()
grid_table(doc,
    ["Component",        "Length", "Description",                    "Example"],
    [["Institution Code","4 chars","Unique identifier for the bank",  "ABCD"],
     ["Country Code",    "2 chars","ISO 3166-1 alpha-2 country code", "GB"],
     ["Location Code",   "2 chars","City / regional identifier",      "2L"],
     ["Branch Code",     "3 chars","Optional — specific branch/dept", "XXX"]],
    col_widths=[1.55, 0.9, 2.8, 1.22])

heading(doc, "Connectivity Model Comparison", SWIFT_BLUE, 11, space_before=4)
grid_table(doc,
    ["Model",             "Best For",                   "Infrastructure",              "Setup Time", "HSM Required"],
    [["SWIFTNet Link\n(Direct/On-Premise)",
                          "High-volume institutions",   "Dedicated servers + HSM",    "8–16 wks",   "Yes"],
     ["Alliance Lite2\n(Cloud)",
                          "Low-to-mid volume",          "SWIFT-managed cloud",         "3–6 wks",    "No"],
     ["Service Bureau",   "Minimal IT footprint",       "Fully outsourced to partner", "2–4 wks",    "No"],
     ["SWIFT API\n(Connectivity+)",
                          "API-first integrations",     "REST/OAuth2 over HTTPS",      "4–8 wks",    "No"]],
    col_widths=[1.35, 1.45, 1.65, 1.0, 1.02])
divider(doc)

# ─── SECTION 02 – PREREQUISITES ──────────────────────────────────────────────
doc.add_paragraph()
banner(doc, "02  |  PREREQUISITES — BEFORE YOU BEGIN")
heading(doc, "Organisational Requirements", SWIFT_BLUE, 11, space_before=4)
grid_table(doc,
    ["#", "Requirement",                                                              "Owner",             "Status"],
    [["1", "Regulated financial institution licence (bank / payment / FI)",           "Legal & Compliance", "☐"],
     ["2", "Legal Entity Identifier (LEI) — registered with an LOI",                 "Compliance Officer", "☐"],
     ["3", "Board resolution authorising SWIFT membership",                           "Board Secretary",    "☐"],
     ["4", "Appointed SWIFT Security Officer (SSO) — CSP mandate",                   "CISO",               "☐"],
     ["5", "Appointed SWIFT Relationship Manager (internal coordinator)",             "Operations",         "☐"],
     ["6", "Budget approved (annual SWIFT fees + setup costs)",                       "Finance",            "☐"]],
    col_widths=[0.28, 3.39, 1.6, 0.72])

heading(doc, "Technical Requirements", SWIFT_BLUE, 11, space_before=4)
grid_table(doc,
    ["#", "Requirement",                                                              "Owner",          "Status"],
    [["1", "Connectivity model decision (Cloud / On-Premise / Bureau)",               "IT Architecture","☐"],
     ["2", "Outbound HTTPS / TLS 1.2+ connectivity to SWIFT endpoints",               "Network/Sec",    "☐"],
     ["3", "SWIFT PKI certificate procurement plan",                                  "IT Security",    "☐"],
     ["4", "Dedicated SWIFT operator accounts and workstations",                      "IT Operations",  "☐"],
     ["5", "Core banking system API / adapter identified for SWIFT integration",      "IT Architecture","☐"],
     ["6", "HSM procurement (on-premise only)",                                       "IT Infra",       "☐"]],
    col_widths=[0.28, 3.39, 1.32, 0.72])
divider(doc)

# ─── SECTION 03 – ONBOARDING PHASES OVERVIEW ─────────────────────────────────
doc.add_paragraph()
banner(doc, "03  |  ONBOARDING PHASES — STEP-BY-STEP")
body(doc, "Seven phases from eligibility through to live API-based message retrieval.", italic=True)
doc.add_paragraph()

flow_row(doc, [
    ("Ph.1\nMembership\nApplication",    SWIFT_BLUE),
    ("Ph.2\nBIC\nRegistration",          SWIFT_LIGHT),
    ("Ph.3\nCloud Setup\n& Certificates",SWIFT_CLOUD),
    ("Ph.4\nSecurity\n& CSP",            SWIFT_GOLD),
    ("Ph.5\nApp\nSetup",                 SWIFT_GREEN),
    ("Ph.6\nAPI &\nMessages",            SWIFT_TEAL),
    ("Ph.7\nGo-Live\n& Ops",             SWIFT_RED),
])

# ── PHASE 1 ───────────────────────────────────────────────────────────────────
heading(doc, "PHASE 1 — SWIFT Membership Application", SWIFT_BLUE, 12, space_before=10)
divider(doc, '003372', '6')

step_box(doc, "1.1", "Create Account on SWIFT Portal",
    "Navigate to the SWIFT portal and register using your organisation's official email domain.\n"
    "  Portal: swift.com/myswift\n"
    "  • Complete company profile (legal name, country, BIC request intent)\n"
    "  • Nominate a primary contact as SWIFT Relationship Manager\n"
    "  • Portal activation email arrives within 1–2 business days", SWIFT_BLUE)

step_box(doc, "1.2", "Download & Complete Membership Application Pack",
    "From the portal: Documents › New Member Application Pack. Complete all sections:\n"
    "  Section A — Institution Details (legal name, registered address, regulator, LEI)\n"
    "  Section B — Connectivity Model (Cloud / On-Premise / Bureau)\n"
    "  Section C — Service Profile (FIN, FileAct, InterAct, Messaging API)\n"
    "  Section D — Corporate Shareholding / Group Structure\n"
    "  Section E — Authorised Signatory Declaration & Board Resolution", SWIFT_BLUE)

step_box(doc, "1.3", "Upload Supporting Documents",
    "Required attachments in PDF format:\n"
    "  ✓ Certificate of Incorporation\n"
    "  ✓ Regulatory licence (banking / payment institution)\n"
    "  ✓ LEI certificate from a recognised LOI\n"
    "  ✓ Board resolution authorising SWIFT membership\n"
    "  ✓ Audited financial statements (last 2 years)\n"
    "  ✓ Proof of registered office address (utility bill / government letter)", SWIFT_BLUE)

step_box(doc, "1.4", "Submit Application & Pay Admission Fee",
    "Submit via portal. Pay the one-time admission fee (amount varies by membership tier).\n"
    "  • SWIFT acknowledges receipt within 5 business days\n"
    "  • SWIFT membership review period: 4–8 weeks\n"
    "  • You will receive an 'Approval in Principle' letter upon successful review", SWIFT_BLUE)

note_box(doc, "TIMELINE", "Phase 1: 4–8 weeks from submission to membership approval.", SWIFT_BLUE)

# ── PHASE 2 ───────────────────────────────────────────────────────────────────
heading(doc, "PHASE 2 — BIC Registration", SWIFT_LIGHT, 12, space_before=10)
divider(doc, '0070C0', '6')

step_box(doc, "2.1", "Request a New BIC via SWIFT BIC Registry",
    "Log in to the SWIFT BIC Registry: bicregistry.swift.com\n"
    "  • Select 'Request a New BIC'\n"
    "  • Choose format: 8-character (institution) or 11-character (branch)\n"
    "  • Enter institution name, ISO country code, city code\n"
    "  • Portal auto-checks uniqueness against global BIC directory\n"
    "  • Select services to register: FIN / FileAct / InterAct / Messaging API", SWIFT_LIGHT)

step_box(doc, "2.2", "BIC Validation & Allocation",
    "SWIFT validates the proposed BIC code within 5–10 business days:\n"
    "  • If proposed BIC is taken, SWIFT suggests alternatives\n"
    "  • Upon approval, BIC is reserved in the SWIFT BIC Registry\n"
    "  • BIC is published in SWIFT BIC Directory (BICPlus) — globally resolvable\n"
    "  • Confirmation letter sent to authorised signatory", SWIFT_LIGHT)

step_box(doc, "2.3", "Sign SWIFT Membership Agreement",
    "Review and sign the SWIFT Membership Agreement:\n"
    "  • Terms of service, data usage policy, liability clauses\n"
    "  • CSP (Customer Security Programme) obligations\n"
    "  • Sign electronically via portal (qualified e-signature accepted)\n"
    "  • SWIFT Legal team countersigns within 5 business days", SWIFT_LIGHT)

note_box(doc, "TIMELINE", "Phase 2: 1–2 weeks after membership approval.", SWIFT_LIGHT)

# ── PHASE 3 — SWIFT CLOUD SETUP & CERTIFICATES ────────────────────────────────
heading(doc, "PHASE 3 — SWIFT Cloud (Alliance Lite2) Setup & Certificate Management", SWIFT_CLOUD, 12, space_before=10)
divider(doc, '006AB0', '6')

body(doc,
    "Alliance Lite2 (AL2) is SWIFT's cloud-hosted connectivity solution. It requires no on-premise "
    "SWIFT hardware or HSM. It connects your bank application to the SWIFT network via a secure "
    "SWIFT-managed cloud infrastructure over standard HTTPS/TLS.", italic=True)
doc.add_paragraph()

# Architecture diagram
arch_diagram(doc, [
    ("Bank Application  /  Core Banking System",            DARK_GREY,   WHITE),
    ("          ▼   REST/HTTPS  ▼",                          LIGHT_GREY,  SWIFT_BLUE),
    ("SWIFT Connector (installed on bank network)",          SWIFT_BLUE,  WHITE),
    ("          ▼   TLS 1.2+  ▼  mTLS Certificate Auth",   LIGHT_GREY,  SWIFT_CLOUD),
    ("Alliance Lite2 Cloud Gateway (SWIFT-managed)",         SWIFT_CLOUD, WHITE),
    ("          ▼   SWIFTNet IP  ▼",                         LIGHT_GREY,  SWIFT_CLOUD),
    ("SWIFT Network (FIN / FileAct / Messaging API)",        SWIFT_BLUE,  WHITE),
], "Alliance Lite2 Cloud Architecture")

heading(doc, "3A — Activate Alliance Lite2 Cloud Subscription", SWIFT_CLOUD, 10, space_before=8)

step_box(doc, "3A.1", "Order Alliance Lite2 via SWIFT Portal",
    "Log into swift.com/myswift › Store › Alliance Lite2:\n"
    "  • Select subscription tier: AL2 Lite / AL2 Standard / AL2 Plus\n"
    "  • Tier is based on message volume (msgs/day) and services required\n"
    "  • Submit order — SWIFT provisions your cloud tenant within 3–5 business days\n"
    "  • You receive a Welcome Email with your AL2 Tenant ID and activation link", SWIFT_CLOUD)

step_box(doc, "3A.2", "Download SWIFT Connector Software",
    "The SWIFT Connector is a lightweight component installed on your bank's network:\n"
    "  1. Log in to SWIFT Support › Downloads › Alliance Lite2 Connector\n"
    "  2. Download the installer package (Windows Server / Linux .rpm / .deb)\n"
    "  3. Verify the SHA-256 checksum published by SWIFT before installing\n"
    "  4. Install on a dedicated server within your DMZ or secured internal network\n"
    "  System requirements: Windows Server 2019+ or RHEL 8+, 8 GB RAM, 50 GB disk, outbound TCP 443",
    SWIFT_CLOUD)

step_box(doc, "3A.3", "Configure SWIFT Connector",
    "After installation, edit the connector configuration file:\n"
    "  Config file location (Linux): /opt/swift/connector/config/connector.properties\n"
    "  Config file location (Windows): C:\\SWIFT\\Connector\\config\\connector.properties\n\n"
    "  Key parameters to set:\n"
    "    swift.tenant.id          = <your AL2 Tenant ID from Welcome Email>\n"
    "    swift.bic                = <your 8 or 11 character BIC>\n"
    "    swift.institution.id     = <Institution ID from SWIFT portal>\n"
    "    swift.proxy.host         = <proxy hostname if applicable>\n"
    "    swift.proxy.port         = <proxy port>\n"
    "    ssl.keystore.path        = /opt/swift/certs/swift-connector.jks\n"
    "    ssl.keystore.password    = <keystore password>",
    SWIFT_CLOUD)

heading(doc, "3B — PKI Certificate Download & Installation", SWIFT_CLOUD, 10, space_before=10)
body(doc,
    "SWIFT uses mutual TLS (mTLS) authentication. Your bank must obtain and install SWIFT-issued "
    "PKI certificates to establish an authenticated connection to the SWIFT network. "
    "There are three certificate types required:", size=9)
doc.add_paragraph()

grid_table(doc,
    ["Certificate Type",         "Purpose",                                       "Format",  "Validity"],
    [["User Certificate",         "Operator/admin authentication to AL2 portal",  "PEM/P12", "1 year"],
     ["Institution Certificate",  "Institution-level message signing & auth",     "PEM/P12", "3 years"],
     ["Connector Certificate",    "mTLS handshake between Connector and AL2",     "JKS/PEM", "2 years"]],
    col_widths=[1.6, 2.7, 0.85, 0.82])

step_box(doc, "3B.1", "Generate a Certificate Signing Request (CSR)",
    "On your server, generate a CSR using OpenSSL (or Java keytool for JKS format):\n\n"
    "  # Step 1 — Generate a 2048-bit RSA private key\n"
    "  openssl genrsa -out swift-institution.key 2048\n\n"
    "  # Step 2 — Generate the CSR (fill in your institution details when prompted)\n"
    "  openssl req -new -key swift-institution.key \\\n"
    "    -out swift-institution.csr \\\n"
    "    -subj \"/C=GB/O=Your Bank Name/OU=SWIFT/CN=YOURBICXXXX\"\n\n"
    "  # For Connector certificate using Java keytool:\n"
    "  keytool -genkeypair -alias swift-connector \\\n"
    "    -keyalg RSA -keysize 2048 -validity 730 \\\n"
    "    -keystore swift-connector.jks \\\n"
    "    -dname \"CN=YOURBICXXXX, OU=SWIFT, O=Your Bank, C=GB\"",
    SWIFT_CLOUD)

step_box(doc, "3B.2", "Submit CSR to SWIFT Certificate Authority",
    "Submit the CSR (.csr file) to SWIFT CA via the portal:\n"
    "  1. Log in: swift.com/myswift › Security › Certificates › Request Certificate\n"
    "  2. Select certificate type (Institution / User / Connector)\n"
    "  3. Upload the .csr file\n"
    "  4. Confirm institution details and authorised signatory\n"
    "  5. Submit — SWIFT CA signs and issues the certificate within 1–2 business days\n"
    "  6. You receive an email notification when the certificate is ready to download",
    SWIFT_CLOUD)

step_box(doc, "3B.3", "Download Certificate from SWIFT Portal",
    "After SWIFT CA signs your CSR:\n"
    "  1. Log in: swift.com/myswift › Security › Certificates › My Certificates\n"
    "  2. Locate the certificate (status: Issued)\n"
    "  3. Click 'Download' — saves as .pem or .cer file\n"
    "  4. Also download the SWIFT CA Root Certificate and Intermediate CA bundle\n"
    "     (required for full chain validation)\n"
    "  5. Verify certificate details:\n"
    "     openssl x509 -in swift-institution.pem -noout -text\n"
    "     Check: CN matches your BIC, validity dates, issuer = SWIFT CA",
    SWIFT_CLOUD)

step_box(doc, "3B.4", "Install Certificate on SWIFT Connector",
    "Combine the private key and signed certificate into a keystore:\n\n"
    "  # Create PKCS12 bundle from key + signed cert + CA chain\n"
    "  openssl pkcs12 -export \\\n"
    "    -inkey swift-institution.key \\\n"
    "    -in swift-institution.pem \\\n"
    "    -certfile swift-ca-chain.pem \\\n"
    "    -out swift-institution.p12 \\\n"
    "    -name swift-institution\n\n"
    "  # Import into Java KeyStore (if using JKS)\n"
    "  keytool -importkeystore \\\n"
    "    -srckeystore swift-institution.p12 -srcstoretype PKCS12 \\\n"
    "    -destkeystore swift-connector.jks  -deststoretype JKS\n\n"
    "  Update connector.properties with the JKS path and password, then restart the Connector.",
    SWIFT_CLOUD)

step_box(doc, "3B.5", "Certificate Renewal Process",
    "Certificates must be renewed before expiry to avoid service interruption:\n"
    "  • SWIFT sends renewal reminder emails 90 days, 30 days, and 7 days before expiry\n"
    "  • Repeat Steps 3B.1–3B.4 to generate a new CSR and download a new certificate\n"
    "  • Install the new certificate alongside the old one first, then switch over\n"
    "  • Revoke the old certificate via portal: Security › Certificates › Revoke\n"
    "  • Set a calendar reminder — PKI expiry causes immediate connectivity loss",
    SWIFT_CLOUD)

grid_table(doc,
    ["Certificate Alert",        "Action Required",                             "When"],
    [["90 days before expiry",   "Begin CSR generation and submission",         "Planned renewal"],
     ["30 days before expiry",   "Install new certificate in parallel",         "Cut-over testing"],
     ["7 days before expiry",    "Switch to new certificate, revoke old",       "Urgent action"],
     ["Certificate expired",     "Immediate outage — emergency reissue only",  "CRITICAL"]],
    col_widths=[1.6, 3.0, 1.37], hdr_color=SWIFT_RED)

note_box(doc, "TIMELINE", "Phase 3 total: 3–6 weeks (Cloud / AL2). Direct on-premise: 8–12 weeks.", SWIFT_CLOUD)

# ── PHASE 4 — SECURITY & CSP ──────────────────────────────────────────────────
heading(doc, "PHASE 4 — Security & Compliance (CSP / KYC-SA)", SWIFT_GOLD, 12, space_before=10)
divider(doc, 'F0A000', '6')
body(doc, "All SWIFT members must comply with the SWIFT Customer Security Programme (CSP) annually.", italic=True)
doc.add_paragraph()

step_box(doc, "4.1", "Complete CSP Self-Assessment (KYC-SA)",
    "Access the KYC Security Attestation application: swift.com/kycsa\n"
    "  • Complete annual self-assessment against SWIFT CSCF (Customer Security Controls Framework)\n"
    "  • 16 Mandatory controls must be attested as Implemented\n"
    "  • 11 Advisory controls recommended\n"
    "  • Key mandatory controls include: network segmentation, privileged access controls,\n"
    "    software integrity monitoring, multi-factor authentication for SWIFT operators",
    SWIFT_GOLD)

step_box(doc, "4.2", "Independent Assessment (Mandatory)",
    "Engage a SWIFT-approved Independent Assessment Provider (IAP):\n"
    "  • IAP performs an independent audit of your CSP controls\n"
    "  • Upload IAP assessment report to KYC-SA portal\n"
    "  • Non-compliant institutions are flagged to all counterparties (visible on SWIFT portal)\n"
    "  • Find approved IAPs: swift.com/compliance-services/partners",
    SWIFT_GOLD)

step_box(doc, "4.3", "AML & Sanctions Screening Setup",
    "Integrate real-time sanctions screening into the SWIFT message pipeline:\n"
    "  • Connect a sanctions engine (e.g., Fircosoft Continuity, Oracle Banking, WorldCheck)\n"
    "  • Configure screening against: OFAC SDN, EU Consolidated, UN Consolidated, local lists\n"
    "  • Define workflow: auto-block / manual review / escalate to Compliance Officer\n"
    "  • Test with known SDN names before go-live",
    SWIFT_GOLD)

step_box(doc, "4.4", "Audit Logging & Data Retention",
    "Configure message archiving and operator audit trails:\n"
    "  • Archive all inbound and outbound SWIFT messages (minimum 5 years)\n"
    "  • Enable operator action logs in Alliance Lite2 / connector audit logs\n"
    "  • Pipe logs to your SIEM (Splunk, QRadar, Sentinel) for real-time monitoring\n"
    "  • Confirm log retention aligns with local regulatory requirements",
    SWIFT_GOLD)

note_box(doc, "TIMELINE", "Phase 4: 3–6 weeks (CSP assessment + control implementation).", SWIFT_GOLD)

# ── PHASE 5 — APPLICATION SETUP ───────────────────────────────────────────────
heading(doc, "PHASE 5 — Bank Application Setup for SWIFT Connectivity", SWIFT_GREEN, 12, space_before=10)
divider(doc, '007050', '6')
body(doc,
    "This phase covers configuring your bank's core application or middleware layer to communicate "
    "with the SWIFT Connector. Messages flow from your application → SWIFT Connector → SWIFT Cloud → Correspondent.",
    italic=True)
doc.add_paragraph()

step_box(doc, "5.1", "Choose Integration Pattern",
    "Select the integration method between your bank application and SWIFT Connector:\n\n"
    "  OPTION A — File-Based (FileAct / SFTP drop folder)\n"
    "    Application writes XML / MT files to a shared folder\n"
    "    Connector picks up files and routes to SWIFT\n"
    "    Suitable for batch statement retrieval (MT940, CAMT.053)\n\n"
    "  OPTION B — REST API (SWIFT Messaging API)\n"
    "    Application calls SWIFT REST endpoints directly\n"
    "    Real-time / near-real-time message exchange\n"
    "    OAuth 2.0 + mTLS authentication\n\n"
    "  OPTION C — MQ / Message Bus\n"
    "    Application publishes to IBM MQ / ActiveMQ / Kafka\n"
    "    SWIFT Connector subscribes and routes messages",
    SWIFT_GREEN)

step_box(doc, "5.2", "Register Application Credentials on SWIFT Portal",
    "Register your application as an API consumer:\n"
    "  1. Log in: developer.swift.com › My Apps › Create App\n"
    "  2. Enter application name and description\n"
    "  3. Select APIs: Messaging API, Financial Crimes Compliance, Payments\n"
    "  4. Choose environment: Sandbox (testing) or Production\n"
    "  5. Download the generated Client ID and Client Secret\n"
    "     IMPORTANT: Store credentials in a vault (HashiCorp Vault / AWS Secrets Manager)\n"
    "     Never hardcode credentials in application source code",
    SWIFT_GREEN)

step_box(doc, "5.3", "Configure OAuth 2.0 Token Authentication",
    "Your application must obtain a Bearer token before each API call:\n\n"
    "  Token endpoint: https://api.swiftnet.sipn.swift.com/oauth2/v1/token\n\n"
    "  POST /oauth2/v1/token\n"
    "  Content-Type: application/x-www-form-urlencoded\n"
    "  Body:\n"
    "    grant_type=client_credentials\n"
    "    &client_id=<your_client_id>\n"
    "    &client_secret=<your_client_secret>\n"
    "    &scope=swift.preval:read swift.messaging:read\n\n"
    "  Response:\n"
    "    { \"access_token\": \"eyJhbGci...\", \"token_type\": \"Bearer\", \"expires_in\": 3600 }\n\n"
    "  Cache the token and refresh 60 seconds before expiry to avoid call failures",
    SWIFT_GREEN)

step_box(doc, "5.4", "Configure mTLS Client Certificate",
    "All SWIFT API calls require mutual TLS in addition to OAuth:\n"
    "  • Use the Institution Certificate downloaded in Phase 3 (Step 3B.3)\n"
    "  • Configure your HTTP client / API gateway with the client certificate:\n\n"
    "    curl example:\n"
    "    curl --cert swift-institution.pem \\\n"
    "         --key  swift-institution.key \\\n"
    "         --cacert swift-ca-chain.pem \\\n"
    "         -H 'Authorization: Bearer <token>' \\\n"
    "         https://api.swiftnet.sipn.swift.com/swift-apitracker/v2/payments\n\n"
    "  • In Java: configure SSLContext with KeyStore containing the client certificate\n"
    "  • In Python (requests): use cert=('cert.pem','key.pem') parameter",
    SWIFT_GREEN)

step_box(doc, "5.5", "Configure Message Routing & Queue",
    "Set up routing rules for inbound and outbound messages in the connector:\n"
    "  Connector admin URL: https://localhost:9443/connector-admin\n\n"
    "  Outbound routing rules:\n"
    "    MT103  → FIN service → route to correspondent BIC\n"
    "    MT202  → FIN service → route to correspondent BIC\n"
    "    PAIN.001 → FileAct service → route to clearing BIC\n\n"
    "  Inbound message handlers:\n"
    "    MT940  → write to /data/swift/inbound/statements/ folder\n"
    "    CAMT.053 → POST to https://your-bank-api/swift/inbound/camt053\n"
    "    MT103  → POST to https://your-bank-api/swift/inbound/payments",
    SWIFT_GREEN)

step_box(doc, "5.6", "Test Connectivity in SWIFT Sandbox",
    "Before production, validate your integration in the SWIFT Sandbox environment:\n"
    "  Sandbox base URL: https://sandbox.swift.com/swift-apitracker/v2/\n"
    "  1. Use Sandbox Client ID / Secret from developer.swift.com\n"
    "  2. Send a test MT103 payment — verify ACK/NACK response\n"
    "  3. Subscribe to test MT940 statement delivery\n"
    "  4. Verify mTLS handshake succeeds (check TLS handshake logs)\n"
    "  5. Confirm OAuth token refresh logic works across token expiry boundary\n"
    "  6. Load test with 100 messages — check throughput and error handling",
    SWIFT_GREEN)

note_box(doc, "TIMELINE", "Phase 5: 3–6 weeks (development, integration testing, sandbox validation).", SWIFT_GREEN)

# ── PHASE 6 — MT940 & CAMT.053 VIA SWIFT API ──────────────────────────────────
heading(doc, "PHASE 6 — Retrieving MT940 & CAMT.053 Messages via SWIFT", SWIFT_TEAL, 12, space_before=10)
divider(doc, '008B8B', '6')
body(doc,
    "MT940 (Customer Statement Message) and CAMT.053 (Bank-to-Customer Statement ISO 20022) "
    "provide end-of-day account statement data from correspondent banks. This phase covers both "
    "FileAct batch retrieval and the SWIFT Messaging REST API for real-time access.",
    italic=True)
doc.add_paragraph()

heading(doc, "6A — MT940 & CAMT.053 Message Formats", SWIFT_TEAL, 10, space_before=6)
grid_table(doc,
    ["Format",      "Standard",   "Type",     "Encoding",   "Typical Schedule",  "Use Case"],
    [["MT940",      "SWIFT FIN",  "Flat text","Proprietary","End-of-day (EOD)",  "Legacy statement reconciliation"],
     ["MT942",      "SWIFT FIN",  "Flat text","Proprietary","Intraday",          "Intraday liquidity monitoring"],
     ["CAMT.053",   "ISO 20022",  "XML",      "UTF-8",      "EOD / on-demand",   "Rich structured statement data"],
     ["CAMT.052",   "ISO 20022",  "XML",      "UTF-8",      "Intraday",          "Intraday account report"]],
    col_widths=[0.9, 1.0, 0.85, 1.0, 1.3, 1.42])

heading(doc, "6B — Method 1: FileAct Batch Retrieval (Pull)", SWIFT_TEAL, 10, space_before=6)
body(doc, "FileAct is SWIFT's file transfer service. Correspondent banks send statement files to your BIC at agreed times.")

step_box(doc, "6B.1", "Subscribe to Statement Delivery via SWIFT Portal",
    "Configure statement delivery with each correspondent bank:\n"
    "  1. Notify your correspondent bank's SWIFT team of your BIC and required message type\n"
    "     (MT940 or CAMT.053) and delivery schedule (EOD / intraday)\n"
    "  2. Correspondent registers your BIC as statement recipient in their SWIFT configuration\n"
    "  3. On your SWIFT portal: Configure › FileAct › Authorised Senders\n"
    "     Add your correspondents' BICs to the authorised sender list\n"
    "  4. Define local file drop path: /data/swift/inbound/statements/",
    SWIFT_TEAL)

step_box(doc, "6B.2", "Configure FileAct Connector to Receive Files",
    "In connector.properties, configure the FileAct inbound handler:\n\n"
    "  fileact.inbound.enabled          = true\n"
    "  fileact.inbound.drop.folder      = /data/swift/inbound/statements/\n"
    "  fileact.inbound.file.pattern     = *.sta, *.xml\n"
    "  fileact.inbound.notification.url = https://your-bank-api/swift/statements/received\n"
    "  fileact.inbound.archive.folder   = /data/swift/archive/statements/\n\n"
    "  The connector will write received files and POST a notification to your application webhook.",
    SWIFT_TEAL)

step_box(doc, "6B.3", "Parse MT940 Statement in Your Application",
    "MT940 files arrive as plain text (SWIFT FIN format). Parse using a FIN parser library:\n\n"
    "  MT940 structure overview:\n"
    "    :20:  Transaction Reference Number\n"
    "    :25:  Account Identification (BIC + Account Number)\n"
    "    :28C: Statement / Sequence Number\n"
    "    :60F: Opening Balance (D/C, Date, Currency, Amount)\n"
    "    :61:  Statement Line (Value Date, Amount, Transaction Reference)\n"
    "    :86:  Information to Account Owner (narrative / structured info)\n"
    "    :62F: Closing Balance\n"
    "    :64:  Available Balance\n\n"
    "  Python example using mt940 library:\n"
    "    pip install mt940\n"
    "    import mt940\n"
    "    transactions = mt940.models.Transactions()\n"
    "    with open('statement.sta', 'rb') as f:\n"
    "        transactions.parse(f.read())\n"
    "    for txn in transactions.transactions:\n"
    "        print(txn.data['date'], txn.data['amount'])",
    SWIFT_TEAL)

step_box(doc, "6B.4", "Parse CAMT.053 XML Statement",
    "CAMT.053 is ISO 20022 XML. Parse using standard XML libraries:\n\n"
    "  CAMT.053 key XML elements:\n"
    "    <GrpHdr>    — Group Header (message ID, creation date)\n"
    "    <Stmt>      — Statement block\n"
    "      <Id>      — Statement ID\n"
    "      <AcctId>  — Account Identification (IBAN)\n"
    "      <Bal>     — Opening, Closing, Available Balances\n"
    "      <Ntry>    — Entry (transaction)\n"
    "        <Amt>   — Amount and currency\n"
    "        <CdtDbtInd> — CRDT (credit) or DBIT (debit)\n"
    "        <ValDt> — Value date\n"
    "        <NtryDtls><TxDtls><RmtInf> — Remittance / reference info\n\n"
    "  Python example:\n"
    "    import xml.etree.ElementTree as ET\n"
    "    ns = {'camt': 'urn:iso:std:iso:20022:tech:xsd:camt.053.001.08'}\n"
    "    tree = ET.parse('statement.xml')\n"
    "    for entry in tree.findall('.//camt:Ntry', ns):\n"
    "        amt = entry.find('camt:Amt', ns).text\n"
    "        ind = entry.find('camt:CdtDbtInd', ns).text\n"
    "        print(f'{ind}: {amt}')",
    SWIFT_TEAL)

heading(doc, "6C — Method 2: SWIFT Transaction Reporting API (Real-Time Pull)", SWIFT_TEAL, 10, space_before=8)
body(doc, "The SWIFT API enables on-demand retrieval of statements and payment status via REST. No file transfers needed.")

step_box(doc, "6C.1", "Authenticate & Get Bearer Token",
    "Obtain an OAuth 2.0 Bearer token (see Phase 5, Step 5.3).\n"
    "Token scope for statement retrieval:\n"
    "  scope=swift.messaging:read swift.accounts:read",
    SWIFT_TEAL)

step_box(doc, "6C.2", "Request Account Statements via SWIFT Messaging API",
    "Use the SWIFT Messaging API to retrieve statements on demand:\n\n"
    "  Endpoint: GET /swift-apitracker/v2/statements\n"
    "  Base URL:  https://api.swiftnet.sipn.swift.com\n\n"
    "  Query parameters:\n"
    "    ?uetr=<unique-end-to-end-transaction-reference>   (for specific payment)\n"
    "    ?account_servicer_bic=CORRESPONDENT_BIC\n"
    "    ?statement_date=2026-04-22\n"
    "    ?message_type=MT940   OR   message_type=CAMT053\n\n"
    "  Full cURL request:\n"
    "  curl -X GET \\\n"
    "    'https://api.swiftnet.sipn.swift.com/swift-apitracker/v2/statements\n"
    "     ?account_servicer_bic=BANKGB2LXXX&statement_date=2026-04-22&message_type=MT940' \\\n"
    "    -H 'Authorization: Bearer <access_token>' \\\n"
    "    --cert swift-institution.pem --key swift-institution.key",
    SWIFT_TEAL)

step_box(doc, "6C.3", "API Response Structure — MT940",
    "Successful API response (HTTP 200) for MT940:\n\n"
    "  {\n"
    "    \"statement_id\": \"STMT-20260422-001\",\n"
    "    \"message_type\": \"MT940\",\n"
    "    \"account_bic\":  \"YOURBICXXXX\",\n"
    "    \"statement_date\": \"2026-04-22\",\n"
    "    \"content_format\": \"FIN\",\n"
    "    \"content\": \":20:STMT20260422001\\n:25:YOURBICXXXX/GB12YOURBANK001\\n\n"
    "      :28C:00001/001\\n:60F:C260421GBP1250000,00\\n\n"
    "      :61:2604220422D15000,00NCHKREF123//COUNTERPARTY REF\\n\n"
    "      :86:/BENEF/VENDOR PAYMENT APRIL\\n\n"
    "      :62F:C260422GBP1235000,00\\n\"\n"
    "  }\n\n"
    "  Parse the 'content' field using your MT940 parser (see Step 6B.3).",
    SWIFT_TEAL)

step_box(doc, "6C.4", "API Response Structure — CAMT.053",
    "Successful API response for CAMT.053 (HTTP 200):\n\n"
    "  {\n"
    "    \"statement_id\": \"CAMT-20260422-001\",\n"
    "    \"message_type\": \"CAMT053\",\n"
    "    \"account_bic\":  \"YOURBICXXXX\",\n"
    "    \"statement_date\": \"2026-04-22\",\n"
    "    \"content_format\": \"XML\",\n"
    "    \"content\": \"<?xml version='1.0' encoding='UTF-8'?>\n"
    "      <Document xmlns='urn:iso:std:iso:20022:tech:xsd:camt.053.001.08'>\n"
    "        <BkToCstmrStmt><GrpHdr><MsgId>CAMT20260422</MsgId>...\n"
    "        <Stmt><Ntry><Amt Ccy='GBP'>15000.00</Amt>\n"
    "        <CdtDbtInd>DBIT</CdtDbtInd>...</Ntry></Stmt></BkToCstmrStmt>\n"
    "      </Document>\"\n"
    "  }\n\n"
    "  Parse the 'content' XML field using your CAMT.053 parser (see Step 6B.4).",
    SWIFT_TEAL)

step_box(doc, "6C.5", "Automate Daily Statement Retrieval",
    "Implement a scheduled job to pull statements automatically:\n\n"
    "  Recommended schedule:\n"
    "    EOD MT940:    Daily at 22:00 UTC (after SWIFT cut-off)\n"
    "    EOD CAMT.053: Daily at 22:15 UTC\n"
    "    Intraday MT942/CAMT.052: Every 4 hours (08:00, 12:00, 16:00 UTC)\n\n"
    "  Implementation checklist:\n"
    "    ✓ Token refresh logic — re-authenticate if token expires\n"
    "    ✓ Idempotency — check statement_id to avoid duplicate processing\n"
    "    ✓ Retry logic — exponential backoff on HTTP 429 (rate limit) or 503\n"
    "    ✓ Dead-letter queue — alert ops team on repeated failures\n"
    "    ✓ Store raw message + parsed data in your reconciliation database\n"
    "    ✓ Post to core banking ledger after successful parse and validation",
    SWIFT_TEAL)

heading(doc, "6D — API Error Codes & Troubleshooting", SWIFT_TEAL, 10, space_before=6)
grid_table(doc,
    ["HTTP Code", "SWIFT Error",          "Cause",                                    "Action"],
    [["200",      "Success",              "Statement retrieved successfully",          "Parse and process content"],
     ["400",      "INVALID_REQUEST",      "Malformed query parameters",               "Check BIC format and date"],
     ["401",      "UNAUTHORIZED",         "Expired or missing Bearer token",          "Re-authenticate, refresh token"],
     ["403",      "FORBIDDEN",            "Client cert mismatch or scope missing",    "Check mTLS cert and OAuth scope"],
     ["404",      "NOT_FOUND",            "No statement for the requested date/BIC",  "Confirm delivery with correspondent"],
     ["429",      "RATE_LIMIT_EXCEEDED",  "Too many requests",                        "Exponential backoff, retry after"],
     ["503",      "SERVICE_UNAVAILABLE",  "SWIFT API maintenance window",             "Retry after 5–10 minutes"]],
    col_widths=[0.75, 1.65, 2.3, 1.77])

note_box(doc, "TIMELINE", "Phase 6: 2–4 weeks (development, sandbox testing, correspondent configuration).", SWIFT_TEAL)

# ── PHASE 7 — GO-LIVE ─────────────────────────────────────────────────────────
heading(doc, "PHASE 7 — Go-Live & Live Operations", SWIFT_RED, 12, space_before=10)
divider(doc, 'C00000', '6')

step_box(doc, "7.1", "Submit Go-Live Request",
    "On SWIFT portal: My Account › Go-Live Request. Attach:\n"
    "  ✓ UAT sign-off report (all test scenarios passed)\n"
    "  ✓ CSP attestation confirmation\n"
    "  ✓ Sanctions screening certification\n"
    "  ✓ Production certificate installation confirmation\n"
    "  SWIFT activates production BIC within 2–3 business days",
    SWIFT_RED)

step_box(doc, "7.2", "Switch Connector to Production Environment",
    "In connector.properties, change from Sandbox to Production:\n\n"
    "  swift.environment = PRODUCTION\n"
    "  swift.api.base.url = https://api.swiftnet.sipn.swift.com\n"
    "  ssl.keystore.path  = /opt/swift/certs/production-connector.jks\n\n"
    "  Restart SWIFT Connector service\n"
    "  Validate production mTLS handshake: check connector logs for 'TLS handshake OK'\n"
    "  Send a live MT202 echo to SWIFT's own BIC (SWHQBEBB) to confirm production link",
    SWIFT_RED)

step_box(doc, "7.3", "Notify Correspondents & Update BIC Directory",
    "Communicate your live BIC to all counterparties:\n"
    "  • Correspondents: formal email from authorised signatory with live BIC\n"
    "  • SWIFT BIC Directory: auto-updated upon go-live activation\n"
    "  • Your bank's public website: update SWIFT/BIC code in contact details\n"
    "  • Central bank / local clearing house: submit BIC registration form\n"
    "  • IBAN registry: update if applicable to your jurisdiction",
    SWIFT_RED)

step_box(doc, "7.4", "Day-1 Operations Checklist",
    "Verify before processing first live payment:\n"
    "  ✓ Production Bearer token obtained and mTLS tested\n"
    "  ✓ MT940/CAMT.053 scheduled jobs configured and tested\n"
    "  ✓ Sanctions screening active and monitoring in production\n"
    "  ✓ Audit logging active and flowing to SIEM\n"
    "  ✓ Incident response runbook distributed to Ops and IT teams\n"
    "  ✓ SWIFT 24×7 support line registered (support.swift.com)\n"
    "  ✓ Nostro accounts funded with correspondent banks\n"
    "  ✓ Cut-off times communicated to treasury and business units\n"
    "  ✓ Certificate expiry alerts configured in monitoring system",
    SWIFT_RED)

note_box(doc, "TIMELINE", "Phase 7: 1–2 weeks. End-to-end total: 4–7 months from application to live operations.", SWIFT_RED)
divider(doc)

# ─── SECTION 04 — MASTER TIMELINE ────────────────────────────────────────────
doc.add_paragraph()
banner(doc, "04  |  MASTER ONBOARDING TIMELINE")
grid_table(doc,
    ["Phase", "Activity",                             "Duration",    "Key Milestone"],
    [["Ph. 1", "SWIFT Membership Application",         "4–8 weeks",   "Approval in Principle Letter"],
     ["Ph. 2", "BIC Registration & Agreement",         "1–2 weeks",   "BIC Published in Directory"],
     ["Ph. 3", "Cloud Setup & Certificate Management", "3–6 weeks",   "mTLS Connectivity Confirmed"],
     ["Ph. 4", "Security & CSP Compliance",            "3–6 weeks",   "CSP Attestation Submitted"],
     ["Ph. 5", "Application Setup & Sandbox Testing",  "3–6 weeks",   "App-to-SWIFT Flow Validated"],
     ["Ph. 6", "MT940 / CAMT.053 API Integration",     "2–4 weeks",   "Statements Retrieved & Parsed"],
     ["Ph. 7", "Go-Live & Live Operations",            "1–2 weeks",   "First Live Message Processed"],
     ["TOTAL", "End-to-End Onboarding Journey",        "4–7 months",  "Fully Operational on SWIFT"]],
    col_widths=[0.72, 2.55, 1.1, 2.1])
divider(doc)

# ─── SECTION 05 — ROLES & RESPONSIBILITIES ───────────────────────────────────
doc.add_paragraph()
banner(doc, "05  |  ROLES & RESPONSIBILITIES")
grid_table(doc,
    ["Role",                        "Key Responsibilities",                                                 "Phase"],
    [["SWIFT Security Officer",      "CSP attestation, PKI cert oversight, incident reporting",             "All"],
     ["SWIFT Relationship Mgr",     "Portal submissions, SWIFT liaison, membership coordination",           "1–2, 7"],
     ["IT / Network Engineer",       "Connector install, certificate management, firewall, HSM",            "3, 5"],
     ["Application Developer",      "OAuth2 integration, API calls, MT940/CAMT.053 parsers",               "5–6"],
     ["Compliance Officer",         "AML/sanctions config, regulatory notifications, LEI upkeep",          "4, 7"],
     ["Operations Manager",         "UAT sign-off, operator training, cut-off time management",            "5–7"],
     ["Finance / Treasury",          "SWIFT fee payments, nostro funding, liquidity monitoring",            "1, 7"],
     ["Legal / Board",               "Board resolution, membership agreement sign-off, LEI registration",  "1"]],
    col_widths=[1.6, 3.6, 1.27])
divider(doc)

# ─── SECTION 06 — KEY REFERENCES ──────────────────────────────────────────────
doc.add_paragraph()
banner(doc, "06  |  KEY CONTACTS & REFERENCES")
grid_table(doc,
    ["Resource",                    "Purpose",                                              "Portal / Access"],
    [["SWIFT Portal (MySwift)",      "Membership, BIC request, CSP attestation",            "swift.com/myswift"],
     ["SWIFT BIC Registry",          "BIC registration and directory lookup",               "bicregistry.swift.com"],
     ["SWIFT Developer Portal",      "API documentation, sandbox, app registration",        "developer.swift.com"],
     ["SWIFT API Reference",         "Messaging API, OAuth2, endpoint specs",               "developer.swift.com/apis"],
     ["SWIFT Customer Support",      "24×7 technical and operational helpdesk",             "support.swift.com"],
     ["SWIFT KYC-SA (CSP)",         "Security attestation submission",                     "swift.com/kycsa"],
     ["SWIFT Pilot / Sandbox",       "Testing environment for messages and APIs",           "pilot.swift.com"],
     ["SWIFT CSCF Framework",        "Customer Security Controls Framework documentation",  "swift.com/myswift/cscf"],
     ["ISO 20022 Catalogue",         "CAMT.053 / PACS.008 XML schemas and specs",          "iso20022.org/catalogue"]],
    col_widths=[1.7, 2.7, 2.07])
divider(doc)

# ─── SECTION 07 — ONGOING COMPLIANCE ─────────────────────────────────────────
doc.add_paragraph()
banner(doc, "07  |  ONGOING COMPLIANCE & OPERATIONAL OBLIGATIONS")
grid_table(doc,
    ["Activity",                              "Frequency",   "Owner",               "Deadline"],
    [["CSP / KYC-SA attestation renewal",     "Annual",      "SWIFT Security Officer","31 December"],
     ["Independent CSP assessment",           "Annual",      "SWIFT Security Officer","31 December"],
     ["PKI certificate renewal (all types)",  "Per expiry",  "IT Security",          "90 days prior"],
     ["SWIFT software / connector upgrade",   "Per release", "IT Operations",        "Per SWIFT schedule"],
     ["BIC directory details review",         "Annual",      "SWIFT Relationship Mgr","January"],
     ["OAuth2 client secret rotation",        "Semi-annual", "IT Security",          "June & December"],
     ["Sanctions list update validation",     "Quarterly",   "Compliance Officer",   "Quarterly"],
     ["Operator access review",               "Semi-annual", "IT Security",          "June & December"],
     ["API statement retrieval audit",        "Monthly",     "Operations Manager",   "Monthly"],
     ["SWIFT fee payment",                    "Annual",      "Finance",              "Per invoice"]],
    col_widths=[2.3, 1.0, 1.65, 1.52])
divider(doc)

# ─── FOOTER ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_tbl = doc.add_table(rows=1, cols=1)
fc = footer_tbl.cell(0, 0)
set_cell_bg(fc, SWIFT_BLUE)
set_no_border(fc)
fp = fc.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(7)
fp.paragraph_format.space_after  = Pt(7)
fr = fp.add_run(
    "CONFIDENTIAL — FOR INTERNAL USE ONLY  |  SWIFT Security Officer  |  "
    "Version 2.0  |  April 2026  |  Subject to annual review"
)
fr.font.color.rgb = MID_GREY
fr.font.size      = Pt(7.5)
fr.font.name      = 'Calibri'

# ─── SAVE ─────────────────────────────────────────────────────────────────────
output_path = r"c:\Subash\Learning\claude\project\Research\SWIFT\Swift-Onboarding-Steps.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
