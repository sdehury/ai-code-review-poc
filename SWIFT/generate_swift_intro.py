from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ──────────────────────────────────────────────────────────────────
SWIFT_BLUE   = RGBColor(0x00, 0x33, 0x72)
SWIFT_LIGHT  = RGBColor(0x00, 0x70, 0xC0)
SWIFT_GOLD   = RGBColor(0xF0, 0xA0, 0x00)
SWIFT_GREEN  = RGBColor(0x00, 0x7A, 0x4C)
SWIFT_TEAL   = RGBColor(0x00, 0x8B, 0x8B)
SWIFT_RED    = RGBColor(0xC0, 0x00, 0x00)
PURPLE       = RGBColor(0x5B, 0x2C, 0x8D)
ORANGE       = RGBColor(0xD4, 0x6A, 0x00)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY   = RGBColor(0xF4, 0xF6, 0xF8)
MID_GREY     = RGBColor(0xBB, 0xBB, 0xBB)
DARK_GREY    = RGBColor(0x22, 0x22, 0x22)
SLATE        = RGBColor(0x2C, 0x3E, 0x50)
CREAM        = RGBColor(0xFF, 0xFD, 0xF0)
TEAL_DARK    = RGBColor(0x00, 0x5F, 0x73)
AMBER        = RGBColor(0xE9, 0xC4, 0x6A)
SAND         = RGBColor(0xF4, 0xA2, 0x61)

def rgb_hex(c): return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

def shade(cell, c):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    for x in p.findall(qn('w:shd')): p.remove(x)
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),rgb_hex(c))
    p.append(s)

def no_border(cell):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        t = OxmlElement(f'w:{e}')
        t.set(qn('w:val'),'none'); t.set(qn('w:sz'),'0')
        t.set(qn('w:space'),'0'); t.set(qn('w:color'),'auto')
        b.append(t)
    p.append(b)

def thin_border(cell, col='CCCCCC'):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        t = OxmlElement(f'w:{e}')
        t.set(qn('w:val'),'single'); t.set(qn('w:sz'),'4')
        t.set(qn('w:space'),'0'); t.set(qn('w:color'),col)
        b.append(t)
    p.append(b)

def add_run(para, text, color=DARK_GREY, size=10, bold=False, italic=False, name='Calibri'):
    r = para.add_run(text)
    r.font.color.rgb = color; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic; r.font.name = name
    return r

def para(doc, space_before=3, space_after=3, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = align
    return p

def h1(doc, text):
    p = para(doc, 16, 4)
    add_run(p, text, SWIFT_BLUE, 18, bold=True)
    return p

def h2(doc, text, color=SWIFT_BLUE):
    p = para(doc, 12, 4)
    add_run(p, text, color, 13, bold=True)
    return p

def h3(doc, text, color=SWIFT_TEAL):
    p = para(doc, 8, 3)
    add_run(p, text, color, 11, bold=True)
    return p

def body_p(doc, text, size=10, italic=False, color=DARK_GREY, indent=0):
    p = para(doc, 3, 4)
    if indent: p.paragraph_format.left_indent = Pt(indent)
    add_run(p, text, color, size, italic=italic)
    return p

def bullet(doc, text, level=0, color=DARK_GREY):
    p = para(doc, 2, 2)
    indent = 14 + level * 14
    p.paragraph_format.left_indent   = Pt(indent)
    p.paragraph_format.first_line_indent = Pt(-10)
    dot = "●" if level == 0 else "○"
    add_run(p, f"{dot}  ", SWIFT_BLUE if level==0 else SWIFT_TEAL, 9, bold=True)
    add_run(p, text, color, 9.5)
    return p

def divider(doc, color='C0C0C0', sz='6'):
    p = doc.add_paragraph()
    pp = p._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),sz)
    bt.set(qn('w:space'),'1'); bt.set(qn('w:color'),color)
    bd.append(bt); pp.append(bd)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after  = Pt(pts)

def banner(doc, text, sub=None, bg=SWIFT_BLUE, fg=WHITE, sub_fg=None, ts=14, ss=9.5):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = tbl.cell(0,0); shade(c, bg); no_border(c)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2 if sub else 10)
    add_run(p, text, fg, ts, bold=True)
    if sub:
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(10)
        add_run(p2, sub, sub_fg or AMBER, ss, italic=True)
    spacer(doc, 4)

def section_banner(doc, num, title, subtitle=None, bg=SWIFT_BLUE):
    tbl = doc.add_table(rows=1, cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    nc = tbl.cell(0,0); tc = tbl.cell(0,1)
    nc.width = Inches(0.52); tc.width = Inches(5.95)
    shade(nc, SWIFT_GOLD); shade(tc, bg)
    no_border(nc); no_border(tc)
    np = nc.paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np.paragraph_format.space_before = Pt(8)
    np.paragraph_format.space_after  = Pt(8)
    add_run(np, num, SWIFT_BLUE, 13, bold=True)
    tp = tc.paragraphs[0]
    tp.paragraph_format.left_indent  = Pt(10)
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after  = Pt(2 if subtitle else 6)
    add_run(tp, title, WHITE, 12.5, bold=True)
    if subtitle:
        tp2 = tc.add_paragraph()
        tp2.paragraph_format.left_indent = Pt(10)
        tp2.paragraph_format.space_after = Pt(6)
        add_run(tp2, subtitle, AMBER, 8.5, italic=True)
    spacer(doc, 4)

def highlight_box(doc, label, text, bg=LIGHT_GREY, label_bg=SWIFT_BLUE, label_fg=WHITE):
    tbl = doc.add_table(rows=1, cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = tbl.cell(0,0); rc = tbl.cell(0,1)
    lc.width = Inches(1.0); rc.width = Inches(5.47)
    shade(lc, label_bg); shade(rc, bg)
    no_border(lc); no_border(rc)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(6)
    lp.paragraph_format.space_after  = Pt(6)
    add_run(lp, label, label_fg, 8.5, bold=True)
    rp = rc.paragraphs[0]
    rp.paragraph_format.left_indent  = Pt(8)
    rp.paragraph_format.space_before = Pt(6)
    rp.paragraph_format.space_after  = Pt(6)
    add_run(rp, text, DARK_GREY, 9.5)
    spacer(doc, 3)

def stat_row(doc, stats):
    """Row of statistic cards  [(value, label, bg_color), ...]"""
    n = len(stats)
    tbl = doc.add_table(rows=1, cols=n); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,(val,lbl,bg) in enumerate(stats):
        c = tbl.cell(0,i); shade(c,bg); no_border(c)
        c.width = Inches(6.47/n)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        add_run(p, val, WHITE, 18, bold=True)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(8)
        add_run(p2, lbl, AMBER, 8, italic=True)
    spacer(doc, 6)

def flow_diagram(doc, nodes, title=None):
    """
    nodes = list of (label, bg, fg)  — rendered as vertical flow with arrows
    """
    if title:
        p = para(doc, 6, 2, WD_ALIGN_PARAGRAPH.CENTER)
        add_run(p, title, SWIFT_BLUE, 9.5, bold=True, italic=True)
    for i,(label,bg,fg) in enumerate(nodes):
        tbl = doc.add_table(rows=1,cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        c = tbl.cell(0,0); shade(c,bg); no_border(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        add_run(p, label, fg, 9, bold=True)
        if i < len(nodes)-1:
            arr = para(doc, 0, 0, WD_ALIGN_PARAGRAPH.CENTER)
            add_run(arr, "▼", SWIFT_GOLD, 11, bold=True)
    spacer(doc, 4)

def horizontal_flow(doc, nodes, arrow_color=SWIFT_GOLD, width=6.47):
    """nodes = [(label, bg, fg), ...]  — horizontal flow with arrows between"""
    cols = len(nodes)*2 - 1
    tbl  = doc.add_table(rows=1, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,(label,bg,fg) in enumerate(nodes):
        ci = i*2
        c  = tbl.cell(0,ci)
        shade(c,bg); no_border(c)
        c.width = Inches((width - (len(nodes)-1)*0.25) / len(nodes))
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after  = Pt(7)
        add_run(p, label, fg, 8.5, bold=True)
        if i < len(nodes)-1:
            ac  = tbl.cell(0, ci+1)
            shade(ac, WHITE); no_border(ac)
            ac.width = Inches(0.25)
            ap = ac.paragraphs[0]
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ap.paragraph_format.space_before = Pt(7)
            add_run(ap, "▶", arrow_color, 11, bold=True)
    spacer(doc, 4)

def two_col_table(doc, headers, rows, col_widths=None, hc=SWIFT_BLUE):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = tbl.rows[0].cells[i]; shade(c,hc)
        thin_border(c, rgb_hex(hc))
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        add_run(p, h, WHITE, 9, bold=True)
    for ri,row in enumerate(rows):
        bg = WHITE if ri%2==0 else LIGHT_GREY
        for ci,val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]; shade(c,bg)
            thin_border(c)
            p = c.paragraphs[0]
            p.paragraph_format.left_indent  = Pt(4)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            add_run(p, str(val), DARK_GREY, 8.5)
    if col_widths:
        for row in tbl.rows:
            for ci,w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    spacer(doc, 4)

def timeline_row(doc, items):
    """items = [(year, event, color), ...]"""
    n   = len(items)
    tbl = doc.add_table(rows=2, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,(year,event,color) in enumerate(items):
        yc = tbl.cell(0,i); ec = tbl.cell(1,i)
        shade(yc, color); shade(ec, LIGHT_GREY)
        no_border(yc); no_border(ec)
        yc.width = Inches(6.47/n); ec.width = Inches(6.47/n)
        yp = yc.paragraphs[0]
        yp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        yp.paragraph_format.space_before = Pt(5)
        yp.paragraph_format.space_after  = Pt(5)
        add_run(yp, year, WHITE, 10, bold=True)
        ep = ec.paragraphs[0]
        ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep.paragraph_format.space_before = Pt(4)
        ep.paragraph_format.space_after  = Pt(4)
        add_run(ep, event, DARK_GREY, 7.5)
    spacer(doc, 4)

def comparison_cards(doc, cards):
    """cards = [(title, points, bg, title_bg), ...]"""
    n = len(cards)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,(title,points,bg,tbg) in enumerate(cards):
        c = tbl.cell(0,i); shade(c,bg); no_border(c)
        c.width = Inches(6.47/n)
        tp = c.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after  = Pt(4)
        add_run(tp, title, tbg, 10, bold=True)
        for pt in points:
            pp = c.add_paragraph()
            pp.paragraph_format.left_indent = Pt(6)
            pp.paragraph_format.space_after  = Pt(3)
            add_run(pp, f"• {pt}", DARK_GREY, 8.5)
        c.add_paragraph().paragraph_format.space_after = Pt(4)
    spacer(doc, 4)

def quote_box(doc, text, attribution=None):
    tbl = doc.add_table(rows=1,cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = tbl.cell(0,0); shade(c, CREAM); thin_border(c, rgb_hex(SWIFT_GOLD))
    p = c.paragraphs[0]
    p.paragraph_format.left_indent  = Pt(16)
    p.paragraph_format.right_indent = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, '\u201c', SWIFT_GOLD, 22, bold=True)
    add_run(p, text, SLATE, 10, italic=True)
    add_run(p, '\u201d', SWIFT_GOLD, 22, bold=True)
    if attribution:
        p2 = c.add_paragraph()
        p2.paragraph_format.left_indent = Pt(16)
        p2.paragraph_format.space_after = Pt(8)
        add_run(p2, f"— {attribution}", SWIFT_TEAL, 8.5, italic=True)
    spacer(doc, 4)

def scenario_box(doc, scenario_title, steps, bg=LIGHT_GREY, title_bg=SWIFT_BLUE):
    """Numbered scenario steps in a styled box."""
    tbl = doc.add_table(rows=1,cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = tbl.cell(0,0); shade(c,bg); thin_border(c, rgb_hex(title_bg))
    tp = c.paragraphs[0]
    tp.paragraph_format.left_indent  = Pt(8)
    tp.paragraph_format.space_before = Pt(7)
    tp.paragraph_format.space_after  = Pt(5)
    add_run(tp, scenario_title, title_bg, 10.5, bold=True)
    for i,(step,detail) in enumerate(steps):
        sp = c.add_paragraph()
        sp.paragraph_format.left_indent = Pt(8)
        sp.paragraph_format.space_after = Pt(3)
        add_run(sp, f"  {i+1:02d}  ", WHITE, 8.5, bold=True, name='Courier New')
        add_run(sp, f"{step}", title_bg, 9, bold=True)
        if detail:
            add_run(sp, f"  —  {detail}", DARK_GREY, 8.5)
    c.add_paragraph().paragraph_format.space_after = Pt(4)
    spacer(doc,3)

def message_box(doc, title, fields, bg=RGBColor(0x1A,0x1A,0x2E), fg=RGBColor(0xA6,0xE3,0xA1), title_color=SWIFT_GOLD):
    """Simulate a SWIFT message in a dark code-style box."""
    tbl = doc.add_table(rows=1,cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = tbl.cell(0,0); shade(c,bg); no_border(c)
    tp = c.paragraphs[0]
    tp.paragraph_format.left_indent  = Pt(8)
    tp.paragraph_format.space_before = Pt(7)
    tp.paragraph_format.space_after  = Pt(4)
    add_run(tp, title, title_color, 9, bold=True, name='Courier New')
    for field in fields:
        fp = c.add_paragraph()
        fp.paragraph_format.left_indent = Pt(8)
        fp.paragraph_format.space_after  = Pt(1)
        add_run(fp, field, fg, 8.5, name='Courier New')
    c.add_paragraph().paragraph_format.space_after = Pt(6)
    spacer(doc,3)

# ═══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.27)
sec.page_height   = Inches(11.69)
sec.left_margin   = Inches(0.85)
sec.right_margin  = Inches(0.85)
sec.top_margin    = Inches(0.65)
sec.bottom_margin = Inches(0.65)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# COVER
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
cover = doc.add_table(rows=1,cols=1)
cc = cover.cell(0,0); shade(cc, SWIFT_BLUE); no_border(cc)

for txt, fg, sz, bold, italic, sb, sa in [
    ("KNOWLEDGE SERIES",                   SWIFT_GOLD,  9,  True,  True,  32, 6 ),
    ("Cross-Border Remittance\n& SWIFT",   WHITE,       26, True,  False, 2,  6 ),
    ("How Money Moves Around the World",   AMBER,       12, False, True,  2,  4 ),
    ("A Beginner-to-Intermediate Guide",   MID_GREY,    9,  False, True,  2,  32),
]:
    p = cc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    add_run(p, txt, fg, sz, bold=bold, italic=italic)

spacer(doc, 6)

# Author / meta strip
meta = doc.add_table(rows=1, cols=4); meta.alignment = WD_TABLE_ALIGNMENT.LEFT
for i,(lbl,val,bg) in enumerate([
    ("Author",      "SWIFT Security Officer",   SLATE),
    ("Series",      "Financial Systems 101",    TEAL_DARK),
    ("Level",       "Beginner → Intermediate",  SWIFT_GREEN),
    ("Published",   "April 2026",               ORANGE),
]):
    c = meta.cell(0,i); shade(c,bg); no_border(c); c.width = Inches(6.47/4)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    add_run(p, lbl+"\n", AMBER, 7, bold=True)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    add_run(p2, val, WHITE, 8.5)
spacer(doc, 8)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 01 — THE GLOBAL MONEY PROBLEM
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
section_banner(doc, "01", "The Global Money Problem",
    "Why sending money across borders was broken — and how the world fixed it", SWIFT_BLUE)

body_p(doc,
    "Imagine you run a small business in London and you've just won a contract with a supplier "
    "in Tokyo. You need to pay ¥5,000,000 (roughly £27,000). Simple, right? You just transfer "
    "the money from your bank account... but how does that money actually get from your London "
    "bank to a Japanese bank on the other side of the world?")
spacer(doc,3)
body_p(doc,
    "This is the fundamental challenge of cross-border remittance — and for most of human history, "
    "it was slow, expensive, unreliable, and riddled with fraud risk.")

spacer(doc,6)
h2(doc, "Before Modern Banking: The Old Ways", SWIFT_RED)
body_p(doc,
    "Before the internet and modern banking networks, international money transfer happened through "
    "a patchwork of fragile, manual systems:")

bullet(doc, "Physical Cash Couriers — Merchants physically carried gold, silver or promissory notes across borders. Risky, slow (weeks to months), and frequently lost to theft.")
bullet(doc, "Bills of Exchange (13th century) — A written order to pay a sum to a named party at a future date. The first 'paper' international payment system, used by the Medici banking family.")
bullet(doc, "Hawala System — An informal trust-based value transfer network originating in South Asia. No physical money moves; brokers (hawaladars) settle via trust and records. Still used in some regions today.")
bullet(doc, "Telegraph / Cable Transfers (1870s) — Western Union began sending money via telegraph wire. The origin of the term 'wire transfer' we still use today.")
bullet(doc, "Telex Systems (1930s–1970s) — Banks used telex machines to instruct each other to move funds. Messages were free-form text — easily forged, with no standard format.")

spacer(doc,4)
h3(doc, "The Core Problem with Telex Banking", SWIFT_RED)
body_p(doc, "By the 1960s, global trade was booming and banks were overwhelmed. Telex had fatal flaws:")

tbl = doc.add_table(rows=1,cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
lc = tbl.cell(0,0); rc = tbl.cell(0,1)
shade(lc, RGBColor(0xFF,0xEB,0xEB)); shade(rc, LIGHT_GREY)
thin_border(lc, 'C00000'); thin_border(rc, 'CCCCCC')
lc.width = Inches(3.0); rc.width = Inches(3.47)
lp = lc.paragraphs[0]
lp.paragraph_format.left_indent  = Pt(8)
lp.paragraph_format.space_before = Pt(6)
add_run(lp, "Problems with Telex\n", SWIFT_RED, 9.5, bold=True)
for pt in ["No standard message format","Easy to forge — no authentication","Errors caused mis-routed payments","Manual re-keying at every bank","No receipt confirmation","Processing took days or weeks"]:
    pp = lc.add_paragraph()
    pp.paragraph_format.left_indent = Pt(8)
    pp.paragraph_format.space_after = Pt(2)
    add_run(pp, f"✗  {pt}", SWIFT_RED, 9)
lc.add_paragraph().paragraph_format.space_after = Pt(6)

rp = rc.paragraphs[0]
rp.paragraph_format.left_indent  = Pt(8)
rp.paragraph_format.space_before = Pt(6)
add_run(rp, "What Banks Needed\n", SWIFT_GREEN, 9.5, bold=True)
for pt in ["A universal message standard","Secure, authenticated messaging","Automated, error-free routing","Instant delivery confirmation","A trusted network all banks could join","A cooperative — not a for-profit monopoly"]:
    pp = rc.add_paragraph()
    pp.paragraph_format.left_indent = Pt(8)
    pp.paragraph_format.space_after = Pt(2)
    add_run(pp, f"✓  {pt}", SWIFT_GREEN, 9)
rc.add_paragraph().paragraph_format.space_after = Pt(6)
spacer(doc,4)

quote_box(doc,
    "In 1973, 239 banks from 15 countries met in Brussels with one shared goal: "
    "build a better way to communicate across borders. The result was SWIFT.",
    "SWIFT History, swift.com")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 02 — WHAT IS SWIFT?
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
divider(doc)
spacer(doc,4)
section_banner(doc,"02","What Is SWIFT?",
    "The financial world's universal language — not a bank, but the network that connects them all",SWIFT_BLUE)

body_p(doc,
    "SWIFT stands for the Society for Worldwide Interbank Financial Telecommunication. "
    "It is a member-owned cooperative headquartered in La Hulpe, Belgium, founded in 1973 "
    "and operational since 1977.")
spacer(doc,3)

highlight_box(doc, "KEY INSIGHT",
    "SWIFT does NOT move money. It moves messages — secure, standardised instructions between banks "
    "that tell them to move money. Think of it as the postal service for financial instructions, "
    "where every letter follows an exact format and is guaranteed to arrive safely.",
    CREAM, SWIFT_GOLD, SWIFT_BLUE)

spacer(doc,4)
stat_row(doc,[
    ("11,000+", "Financial institutions\nin 200+ countries",  SWIFT_BLUE),
    ("44M+",    "Messages sent\nevery single day",            SWIFT_TEAL),
    ("$5T+",    "In payments settled\ndaily through SWIFT",   SWIFT_GREEN),
    ("1977",    "Year SWIFT went\noperationally live",        ORANGE),
])

h2(doc,"What SWIFT Actually Is — and Is NOT", SWIFT_BLUE)
two_col_table(doc,
    ["SWIFT IS...", "SWIFT IS NOT..."],
    [["A secure, encrypted messaging network","A bank — it holds no money"],
     ["A cooperative owned by its member banks","A payment system or clearing house"],
     ["A global standard for financial messages","A government or regulatory body"],
     ["A network connecting 11,000+ institutions","Responsible for settling or clearing funds"],
     ["The creator of BIC codes (SWIFT codes)","Mandatory — banks can choose alternatives"],
     ["Used for payments, securities, trade finance","The only option (SEPA, Fedwire, etc. exist)"]],
    col_widths=[3.23, 3.24])

h2(doc,"SWIFT in History: Key Milestones", SWIFT_BLUE)
timeline_row(doc,[
    ("1973","Founded by\n239 banks",            SWIFT_BLUE),
    ("1977","First live\nmessage sent",          SWIFT_LIGHT),
    ("1987","1,000 member\nbanks reached",       SWIFT_GREEN),
    ("2001","SWIFTNet\nlaunched (IP)",           SWIFT_TEAL),
    ("2015","CSP security\nprogramme launched",  ORANGE),
    ("2022","ISO 20022\nmigration begins",        PURPLE),
    ("2025","Full ISO 20022\nadoption target",    SWIFT_RED),
])

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 03 — THE SWIFT NETWORK ARCHITECTURE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
divider(doc)
spacer(doc,4)
section_banner(doc,"03","The SWIFT Network Architecture",
    "How the plumbing of global finance is actually structured", SWIFT_BLUE)

body_p(doc,
    "SWIFT operates a private, closed-loop network — you cannot send a SWIFT message unless "
    "your institution is a registered SWIFT member with a valid BIC code. This is by design: "
    "the network's value comes from its exclusivity and the trust it creates.")
spacer(doc,3)

h2(doc,"The Three Layers of SWIFT", SWIFT_BLUE)
tbl = doc.add_table(rows=1,cols=3); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i,(title,desc,bg) in enumerate([
    ("LAYER 1\nThe Network\n(SWIFTNet)",
     "The private IP network connecting all member institutions. Messages travel encrypted via TLS. "
     "No public internet. Two geographically redundant Operating Centres (OPCs) — in the Netherlands and USA.",
     SWIFT_BLUE),
    ("LAYER 2\nThe Standards\n(Message Types)",
     "Standardised message formats (MT / MX) ensure every bank speaks the same language. "
     "Field :32A: always means 'value date, currency, amount' — everywhere, every time.",
     SWIFT_TEAL),
    ("LAYER 3\nThe Services\n(FIN / FileAct / InterAct)",
     "FIN: individual message delivery (payments, confirmations).\n"
     "FileAct: bulk file transfer (statements, batch payments).\n"
     "InterAct: real-time request-response (enquiries, status checks).",
     SWIFT_GREEN),
]):
    c = tbl.cell(0,i); shade(c,bg); no_border(c); c.width = Inches(6.47/3)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(4)
    add_run(p, title, WHITE, 9.5, bold=True)
    p2 = c.add_paragraph()
    p2.paragraph_format.left_indent = Pt(6)
    p2.paragraph_format.right_indent = Pt(4)
    p2.paragraph_format.space_after  = Pt(8)
    add_run(p2, desc, AMBER, 8, italic=True)
spacer(doc,6)

h2(doc,"SWIFT BIC Code — The Address of a Bank", SWIFT_BLUE)
body_p(doc,
    "Every institution on the SWIFT network has a unique Bank Identifier Code (BIC) — "
    "sometimes called a SWIFT Code. It is the 'postal address' used to route messages and payments.")
spacer(doc,3)

# BIC anatomy diagram
bic_tbl = doc.add_table(rows=3, cols=4); bic_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
colors  = [SWIFT_BLUE, SWIFT_TEAL, SWIFT_GREEN, ORANGE]
parts   = ["ABCD",   "GB",             "2L",              "XXX"]
labels  = ["Bank\nCode","Country\nCode","Location\nCode", "Branch\nCode"]
descs   = ["4 chars\nIdentifies the\nfinancial institution",
           "2 chars\nISO 3166-1\ncountry code",
           "2 chars\nCity or region\ncode",
           "3 chars\nOptional — HQ if\nomitted (XXX)"]
for ci,(p,l,d,c) in enumerate(zip(parts,labels,descs,colors)):
    shade(bic_tbl.cell(0,ci), c); no_border(bic_tbl.cell(0,ci))
    shade(bic_tbl.cell(1,ci), c); no_border(bic_tbl.cell(1,ci))
    shade(bic_tbl.cell(2,ci), LIGHT_GREY); no_border(bic_tbl.cell(2,ci))
    bic_tbl.cell(0,ci).width = Inches(6.47/4)

    p0 = bic_tbl.cell(0,ci).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(6)
    p0.paragraph_format.space_after  = Pt(6)
    add_run(p0, p, WHITE, 18, bold=True, name='Courier New')

    p1 = bic_tbl.cell(1,ci).paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(5)
    add_run(p1, l, AMBER, 8.5, bold=True)

    p2 = bic_tbl.cell(2,ci).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(4)
    add_run(p2, d, DARK_GREY, 7.5, italic=True)
spacer(doc,3)
p = para(doc, 2, 2, WD_ALIGN_PARAGRAPH.CENTER)
add_run(p, "Example:  ", DARK_GREY, 9)
add_run(p, "ABCD", SWIFT_BLUE, 11, bold=True, name='Courier New')
add_run(p, "GB", SWIFT_TEAL, 11, bold=True, name='Courier New')
add_run(p, "2L", SWIFT_GREEN, 11, bold=True, name='Courier New')
add_run(p, "XXX", ORANGE, 11, bold=True, name='Courier New')
add_run(p, "  =  ABCD Bank, Great Britain, London, Head Office", DARK_GREY, 9, italic=True)
spacer(doc,6)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 04 — HOW A CROSS-BORDER PAYMENT ACTUALLY WORKS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
divider(doc)
spacer(doc,4)
section_banner(doc,"04","How a Cross-Border Payment Works",
    "A step-by-step journey of £27,000 from London to Tokyo — following every hop", SWIFT_BLUE)

body_p(doc,
    "Let's follow our £27,000 payment from a London business to a Tokyo supplier. "
    "This is the most important section — understanding this journey unlocks how the entire "
    "global financial system works.")
spacer(doc,4)

h2(doc,"The Cast of Characters", SWIFT_BLUE)
two_col_table(doc,
    ["Role", "Who They Are", "In Our Example"],
    [["Sender (Originator)",       "The person/company initiating the payment",        "London Business (you)"],
     ["Ordering Bank",             "Your bank — receives your payment instruction",    "Barclays Bank, London (BARCGB22)"],
     ["Correspondent Bank",        "An intermediary bank with accounts at both ends",  "JP Morgan, New York (CHASUSU33)"],
     ["Beneficiary Bank",          "The recipient's bank — holds their account",       "Mizuho Bank, Tokyo (MHCBJPJT)"],
     ["Beneficiary (Receiver)",    "The person/company receiving the money",           "Tokyo Supplier"]],
    col_widths=[1.6, 2.6, 2.27])

spacer(doc,4)
h2(doc,"Why Do We Need Correspondent Banks?", SWIFT_BLUE)
body_p(doc,
    "Barclays in London does not hold a Japanese Yen account at Mizuho in Tokyo — and Mizuho "
    "does not hold a GBP account at Barclays. They have no direct relationship. "
    "This is the core problem of cross-border payments.")
spacer(doc,3)
body_p(doc,
    "The solution: Correspondent Banking — a network of pre-arranged bilateral agreements "
    "where banks hold accounts at each other (called Nostro and Vostro accounts) to facilitate payments on behalf of others.")
spacer(doc,3)

highlight_box(doc, "NOSTRO vs\nVOSTRO",
    "NOSTRO account (from Latin: 'ours') — an account a bank holds at a foreign bank in that bank's currency.\n"
    "VOSTRO account (from Latin: 'yours') — a foreign bank's account held at your bank.\n"
    "Same account, different perspective. Barclays' USD account at JP Morgan = Barclays' Nostro, JP Morgan's Vostro.",
    CREAM, SWIFT_TEAL, WHITE)

spacer(doc,6)
h2(doc, "The Payment Journey — Step by Step", SWIFT_BLUE)
body_p(doc, "Following £27,000 from London to Tokyo:", italic=True)
spacer(doc,4)

# Main journey flow
horizontal_flow(doc, [
    ("You\n(London\nBusiness)", SLATE, WHITE),
    ("Barclays\nLondon\nBARCGB22",  SWIFT_BLUE, WHITE),
    ("JP Morgan\nNew York\nCHASUS33", SWIFT_TEAL, WHITE),
    ("Mizuho\nTokyo\nMHCBJPJT",   SWIFT_GREEN, WHITE),
    ("Tokyo\nSupplier",            ORANGE, WHITE),
])

scenario_box(doc, "STEP-BY-STEP PAYMENT JOURNEY — £27,000  London → Tokyo", [
    ("You instruct your bank",
     "You log in to Barclays online banking and initiate a £27,000 international wire to Tokyo Supplier "
     "at Mizuho Bank (BIC: MHCBJPJT), providing IBAN/account number, amount, and purpose."),
    ("Barclays verifies & sanctions-screens",
     "Barclays checks your account balance, runs the payment against OFAC/UN sanctions lists, "
     "verifies the beneficiary BIC is valid in the SWIFT BIC Directory, and approves the transaction."),
    ("Barclays sends MT103 message via SWIFT",
     "Barclays composes a SWIFT MT103 message (Customer Credit Transfer) and sends it securely "
     "via the SWIFT FIN network to JP Morgan (their USD correspondent). The message contains: "
     "your details, Mizuho's BIC, the converted amount, and the beneficiary's account number."),
    ("JP Morgan receives MT103 & routes onward",
     "JP Morgan debits Barclays' Nostro account held at JP Morgan by the USD equivalent. "
     "JP Morgan then sends a new MT103 message to Mizuho Bank in Tokyo, crediting their account."),
    ("Mizuho receives the message",
     "Mizuho receives the SWIFT MT103, converts the funds to Japanese Yen at the prevailing rate, "
     "and credits your Tokyo supplier's bank account. Mizuho sends an MT910 (Credit Confirmation) "
     "back to JP Morgan."),
    ("Supplier receives funds",
     "The Tokyo supplier sees the credited JPY in their Mizuho account. "
     "A statement entry appears, usually within 1–3 business days of the original instruction."),
    ("Statement messages sent (MT940 / CAMT.053)",
     "At end of day, Mizuho sends an MT940 account statement to Barclays covering all movements "
     "in Barclays' Nostro account. Barclays reconciles this against their internal records."),
], LIGHT_GREY, SWIFT_BLUE)

spacer(doc,4)
highlight_box(doc,"KEY POINT",
    "No money physically moves during this process. What moves are SWIFT messages — instructions "
    "to debit and credit accounts that banks already hold with each other (Nostro/Vostro accounts). "
    "The actual 'settlement' happens by adjusting these pre-funded accounts.",
    CREAM, SWIFT_GOLD, SWIFT_BLUE)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 05 — SWIFT MESSAGE TYPES
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
divider(doc)
spacer(doc,4)
section_banner(doc,"05","SWIFT Message Types — The Language of Banking",
    "Every payment, confirmation and statement has its own message format", SWIFT_BLUE)

body_p(doc,
    "SWIFT messages are organised into Categories (0–9) based on their purpose. "
    "The message type (MT) is a 3-digit code: the first digit is the category, "
    "the next two identify the specific message. Think of it like a filing system for financial instructions.")
spacer(doc,4)

h2(doc,"Message Categories at a Glance", SWIFT_BLUE)
two_col_table(doc,
    ["Category", "Name",                          "Common Messages",           "Used For"],
    [["MT1xx",   "Customer Payments & Cheques",   "MT103, MT102, MT104",       "Cross-border customer payments"],
     ["MT2xx",   "Financial Institution Transfers","MT200, MT202, MT205",      "Bank-to-bank funds transfer"],
     ["MT3xx",   "FX, Money Market & Derivatives","MT300, MT320",              "FX confirmations, interest payments"],
     ["MT4xx",   "Collections & Cash Letters",    "MT400, MT410",              "Documentary collections"],
     ["MT5xx",   "Securities Markets",            "MT515, MT535, MT548",       "Securities settlement, custody"],
     ["MT6xx",   "Precious Metals & Syndications","MT600, MT643",              "Commodity transfers"],
     ["MT7xx",   "Documentary Credits & Guarantees","MT700, MT707, MT750",     "Letters of credit, trade finance"],
     ["MT9xx",   "Cash Management & Statements",  "MT900, MT910, MT940, MT950","Confirmations, statements"]],
    col_widths=[0.75, 1.95, 1.75, 2.02])

spacer(doc,4)
h2(doc,"The Three Most Important Messages", SWIFT_BLUE)

# MT103
h3(doc,"MT103 — Customer Credit Transfer (The Workhorse of Global Payments)")
body_p(doc,
    "The MT103 is the most common SWIFT message in the world. Every time an individual or "
    "business sends an international wire transfer, an MT103 is the message that carries "
    "the instruction across the network.")

message_box(doc, "── SWIFT MT103 SAMPLE MESSAGE ──", [
    "{1:F01BARCGB22AXXX0000000000}",
    "{2:I103MHCBJPJTXXXXN}",
    "{4:",
    ":20:REF20260422001          ← Transaction Reference Number",
    ":23B:CRED                   ← Bank Operation Code (Credit)",
    ":32A:260422GBP27000,00      ← Value Date / Currency / Amount",
    ":50K:/GB12BARC20000012345678  ← Ordering Customer (Sender)",
    "    LONDON BUSINESS LTD",
    "    1 FINANCIAL STREET, LONDON",
    ":57A:MHCBJPJTXXX            ← Account With Institution (Beneficiary Bank)",
    ":59:/1234567890              ← Beneficiary Account & Name",
    "    TOKYO SUPPLIER CO LTD",
    "    3-1 MARUNOUCHI, TOKYO",
    ":70:/INV/2026-04-22/SUPPLY   ← Remittance Information (Invoice ref)",
    ":71A:SHA                    ← Details of Charges (SHA = shared)",
    "-}",
])

body_p(doc,"Field :71A: (Charges) has three options:")
two_col_table(doc,
    ["Code", "Meaning",                                                              "Who Pays Fees"],
    [["SHA", "Shared — most common for international payments",                      "Sender pays sender's bank; beneficiary pays their bank"],
     ["OUR", "Sender pays all fees — full amount guaranteed to arrive",             "Sender covers all correspondent and beneficiary bank fees"],
     ["BEN", "Beneficiary pays all fees — recipient receives net amount after fees","All fees deducted from the transferred amount"]],
    col_widths=[0.55, 3.5, 2.42])

h3(doc,"MT202 — Financial Institution Credit Transfer")
body_p(doc,
    "The MT202 is a bank-to-bank transfer — it moves money between the banks themselves, "
    "not on behalf of a customer. When JP Morgan pays Mizuho on behalf of Barclays, "
    "they use an MT202 (or MT202 COV for cover payments). The MT103 carries the customer "
    "details; the MT202 carries the settlement instruction between the correspondent banks.")

h3(doc,"MT940 / CAMT.053 — Account Statement Messages")
body_p(doc,
    "At the end of each business day, banks send account statements to their correspondents "
    "covering all transactions on Nostro accounts. This is how banks reconcile their books.")

two_col_table(doc,
    ["Format",   "Standard",   "Structure",  "Detail Level",                       "When Used"],
    [["MT940",   "SWIFT FIN",  "Plain text", "Transactions + opening/closing bal", "Legacy systems (still dominant)"],
     ["MT942",   "SWIFT FIN",  "Plain text", "Intraday transaction report",        "Real-time liquidity monitoring"],
     ["CAMT.053","ISO 20022",  "XML",        "Rich structured data, full narrative","Modern ISO 20022 migrations"],
     ["CAMT.052","ISO 20022",  "XML",        "Intraday account report (XML)",      "Intraday ISO 20022 reporting"]],
    col_widths=[0.95, 0.95, 0.9, 2.0, 1.67])

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 06 — CORRESPONDENT BANKING NETWORK
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
divider(doc)
spacer(doc,4)
section_banner(doc,"06","The Correspondent Banking Network",
    "The invisible web of bank-to-bank relationships that makes global finance possible", SWIFT_BLUE)

body_p(doc,
    "No single bank has direct relationships with every other bank in the world. Instead, "
    "banks build a network of bilateral correspondent relationships — a web that, when combined, "
    "connects the entire global financial system.")
spacer(doc,3)

# Correspondent network diagram
h2(doc,"Direct vs Correspondent Routing", SWIFT_BLUE)

tbl = doc.add_table(rows=1,cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
lc = tbl.cell(0,0); rc = tbl.cell(0,1)
lc.width = Inches(3.1); rc.width = Inches(3.37)
shade(lc, RGBColor(0xE8,0xF4,0xF8)); shade(rc, RGBColor(0xE8,0xF8,0xEE))
thin_border(lc,'0070C0'); thin_border(rc,'007A4C')

lh = lc.paragraphs[0]
lh.paragraph_format.left_indent = Pt(6)
lh.paragraph_format.space_before = Pt(6)
add_run(lh, "DIRECT RELATIONSHIP\n", SWIFT_LIGHT, 9.5, bold=True)
add_run(lh, "Bank A holds account at Bank B directly", DARK_GREY, 8.5, italic=True)
for row in [
    ("Bank A (London)","BARCGB22",SWIFT_BLUE),
    ("   Direct SWIFT MT103  ▼","",SWIFT_GOLD),
    ("Bank B (Tokyo)",  "MHCBJPJT",SWIFT_GREEN),
]:
    p = lc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, row[0], row[2], 9, bold=(row[1]!=""))
    if row[1]: add_run(p, f"  [{row[1]}]", DARK_GREY, 7.5)
add_run(lc.add_paragraph(), "✓  Fastest & cheapest route\n✓  Fewer hops = less fees\n✗  Requires pre-existing relationship", SWIFT_GREEN, 8.5)
lc.add_paragraph().paragraph_format.space_after = Pt(6)

rh = rc.paragraphs[0]
rh.paragraph_format.left_indent = Pt(6)
rh.paragraph_format.space_before = Pt(6)
add_run(rh, "CORRESPONDENT ROUTING\n", SWIFT_GREEN, 9.5, bold=True)
add_run(rh, "Via one or more intermediary banks", DARK_GREY, 8.5, italic=True)
for row in [
    ("Bank A (London)",    "BARCGB22", SWIFT_BLUE),
    ("   MT103 via SWIFT  ▼","",       SWIFT_GOLD),
    ("Correspondent Bank", "CHASUSU33",SWIFT_TEAL),
    ("   MT103 via SWIFT  ▼","",       SWIFT_GOLD),
    ("Bank B (Tokyo)",     "MHCBJPJT", SWIFT_GREEN),
]:
    p = rc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, row[0], row[2], 9, bold=(row[1]!=""))
    if row[1]: add_run(p, f"  [{row[1]}]", DARK_GREY, 7.5)
add_run(rc.add_paragraph(), "✓  Works without direct relationship\n✓  Reaches any bank globally\n✗  Slower & more expensive", SWIFT_TEAL, 8.5)
rc.add_paragraph().paragraph_format.space_after = Pt(6)
spacer(doc,4)

h2(doc,"Why Payments Can Take 1–5 Business Days", SWIFT_BLUE)
body_p(doc, "Each 'hop' through a correspondent bank adds time and cost. Here's why:")
bullet(doc, "Sanctions screening at every bank in the chain — each bank runs its own checks")
bullet(doc, "Cut-off times — miss a bank's cut-off and your payment waits until the next business day")
bullet(doc, "Currency conversion — FX rates applied at each conversion point")
bullet(doc, "Compliance reviews — large or unusual payments may be flagged for manual review")
bullet(doc, "Time zones — New York and Tokyo overlap for only 1–2 hours each day")
bullet(doc, "Public holidays — each country has its own, blocking settlement on those days")
spacer(doc,4)

highlight_box(doc,"SWIFT GPI",
    "SWIFT Global Payments Innovation (GPI) was launched in 2017 to address speed and transparency. "
    "GPI payments include a unique end-to-end transaction reference (UETR) that lets all parties "
    "track the payment in real time — like a parcel tracking number for money. "
    "Over 70% of SWIFT cross-border payments now arrive within 30 minutes.",
    CREAM, SWIFT_TEAL, WHITE)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 07 — COSTS, FEES & WHAT AFFECTS YOUR PAYMENT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
divider(doc)
spacer(doc,4)
section_banner(doc,"07","Costs, Fees & What Affects Your Payment",
    "Why cross-border payments cost more than local transfers — and who takes a cut", SWIFT_TEAL)

body_p(doc,
    "International wire transfers are notoriously expensive compared to domestic payments. "
    "Understanding where the fees come from helps you make smarter financial decisions.")
spacer(doc,4)

two_col_table(doc,
    ["Fee Type",                   "Charged By",          "Typical Cost",        "Notes"],
    [["Outgoing wire fee",         "Your bank",           "£10–£40",             "Fixed fee per transaction"],
     ["Correspondent bank fee",    "Each intermediary",   "$10–$30 per hop",     "Can be 2–4 banks in a chain"],
     ["FX conversion spread",      "Your bank / FX desk", "0.5%–3.0% of amount", "Hidden in the exchange rate"],
     ["Lifting fee (incoming)",    "Beneficiary's bank",  "$5–$20",              "Deducted from received amount"],
     ["Compliance/AML review fee", "Any bank in chain",   "Ad hoc",              "For large/unusual payments"],
     ["SWIFT message fee",         "Your bank (SWIFT)",   "< $0.01 per msg",     "Infrastructure cost, passed on"]],
    col_widths=[1.7, 1.6, 1.35, 1.82])

spacer(doc,4)
h2(doc,"Factors That Affect Speed & Cost", SWIFT_TEAL)
two_col_table(doc,
    ["Factor",           "Impact on Payment",                                     "Tip"],
    [["Currency pair",    "USD, EUR, GBP corridors are fastest & cheapest",       "Use major currencies where possible"],
     ["No. of hops",      "Each correspondent bank = more fees + more delays",    "Use banks with direct relationships"],
     ["Payment amount",   "Large payments may trigger enhanced due diligence",    "Pre-notify your bank for large amounts"],
     ["Countries involved","Sanctioned/restricted countries cause delays or blocks","Check OFAC lists before initiating"],
     ["Charge code",      "OUR = predictable full amount; SHA = risk of deductions","Use OUR for guaranteed full delivery"],
     ["Time of day",       "Submit before bank cut-off (often 14:00–15:00 local)", "Submit early in the morning"],
     ["Day of week",       "Friday payments often arrive Monday",                  "Avoid Fridays for urgent payments"]],
    col_widths=[1.3, 2.6, 2.57])

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 08 — SWIFT SECURITY & COMPLIANCE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
divider(doc)
spacer(doc,4)
section_banner(doc,"08","SWIFT Security & Compliance",
    "How the world's most secure financial network protects against fraud and cyber threats", SWIFT_GOLD)

body_p(doc,
    "In February 2016, hackers sent fraudulent SWIFT messages from Bangladesh Bank's systems "
    "and successfully stole $81 million from its Federal Reserve account in New York. "
    "This attack — the largest cyber heist in history — changed SWIFT security forever.")
spacer(doc,4)

highlight_box(doc,"THE\nBANGLADESH\nHEIST 2016",
    "Attackers infiltrated Bangladesh Bank's internal systems, obtained valid SWIFT credentials, "
    "and sent 35 fraudulent MT103 payment instructions to the Federal Reserve Bank of New York "
    "over a weekend. $81M was successfully routed to accounts in the Philippines before the "
    "fraud was discovered. A simple spelling mistake ('fandation' vs 'foundation') flagged "
    "one transfer, saving an additional $850M. This attack led directly to SWIFT's CSP programme.",
    RGBColor(0xFF,0xEB,0xEB), SWIFT_RED, WHITE)

spacer(doc,4)
h2(doc,"SWIFT Customer Security Programme (CSP)", SWIFT_GOLD)
body_p(doc,
    "Launched in 2016 in response to the Bangladesh attack, the CSP is SWIFT's mandatory "
    "security framework for all member institutions. It defines a set of security controls "
    "every SWIFT member must implement and attest to annually.")
spacer(doc,3)

tbl = doc.add_table(rows=1, cols=3); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i,(title,items,bg) in enumerate([
    ("SECURE YOUR\nENVIRONMENT", [
        "Restrict internet access to SWIFT systems",
        "Segment SWIFT network from corporate network",
        "Use dedicated hardware for SWIFT operations",
        "Enforce multi-factor authentication (MFA)",
        "Apply security patches within 30 days",
    ], SWIFT_BLUE),
    ("KNOW & LIMIT\nACCESS", [
        "Restrict user permissions to minimum needed",
        "Review operator accounts every 6 months",
        "Enforce least-privilege access control",
        "Detect and log all privileged access",
        "Physically secure SWIFT workstations",
    ], SWIFT_TEAL),
    ("DETECT &\nRESPOND", [
        "Log all SWIFT operator actions",
        "Monitor transactions 24×7 for anomalies",
        "Integrate with Security Operations Centre",
        "Have an incident response plan ready",
        "Report cyber incidents to SWIFT within 24h",
    ], SWIFT_RED),
]):
    c = tbl.cell(0,i); shade(c, LIGHT_GREY); thin_border(c, rgb_hex(bg)); c.width = Inches(6.47/3)
    tp = c.paragraphs[0]; tp.paragraph_format.space_before = Pt(6); tp.paragraph_format.space_after = Pt(4)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(tp, title, bg, 9.5, bold=True)
    for item in items:
        p = c.add_paragraph()
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"◆  {item}", DARK_GREY, 8)
    c.add_paragraph().paragraph_format.space_after = Pt(6)
spacer(doc,4)

h2(doc,"How SWIFT Messages Are Secured in Transit", SWIFT_GOLD)
flow_diagram(doc,[
    ("1.  Your bank operator creates the SWIFT message", SWIFT_BLUE, WHITE),
    ("2.  Message is digitally signed with PKI private key (authentication)", SWIFT_TEAL, WHITE),
    ("3.  Message is encrypted with TLS for transport (confidentiality)", SWIFT_GREEN, WHITE),
    ("4.  Message enters SWIFTNet — private IP network (no public internet)", SWIFT_GOLD, SWIFT_BLUE),
    ("5.  SWIFT validates digital signature of sender", SLATE, WHITE),
    ("6.  Message delivered to recipient bank — decrypted and signature verified", SWIFT_RED, WHITE),
    ("7.  Recipient bank processes message — ACK (success) or NAK (failure) returned", DARK_GREY, WHITE),
], "SWIFT Message Security Flow")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 09 — SWIFT vs ALTERNATIVES
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
divider(doc)
spacer(doc,4)
section_banner(doc,"09","SWIFT vs Modern Alternatives",
    "Is SWIFT still the best way to move money internationally in 2026?", PURPLE)

body_p(doc,
    "SWIFT has dominated cross-border payments for nearly 50 years — but it's not without "
    "competition. Fintech disruptors, central bank digital currencies, and new networks are "
    "challenging the status quo. Here's how they compare:")
spacer(doc,4)

two_col_table(doc,
    ["Network / Method",   "How It Works",                              "Speed",          "Cost",      "Best For"],
    [["SWIFT (Traditional)","MT messages via correspondent banks",       "1–5 business days","Medium–High","Large, complex institutional payments"],
     ["SWIFT GPI",          "SWIFT with real-time tracking (UETR)",     "< 30 minutes",    "Medium",    "Faster SWIFT payments with transparency"],
     ["SEPA (Europe)",      "EU single payments area — direct bank",    "Same / next day", "Very low",  "EUR payments within the EU/EEA"],
     ["Fedwire (USA)",      "US Federal Reserve real-time gross settlement","Same day",    "Low",       "USD domestic high-value payments"],
     ["Wise (TransferWise)","FX spread + local bank transfer at each end","Minutes–1 day","Very low",  "Consumer/SME low-cost remittances"],
     ["Ripple / XRP",       "Blockchain-based settlement network",      "3–5 seconds",    "Very low",  "Institutional real-time cross-border"],
     ["Central Bank DCCs",  "Digital currencies issued by central banks","Seconds",        "Near zero", "Emerging — future of settlement"],
     ["Hawala (informal)",  "Trust-based broker network, cash offsets", "Minutes–hours",  "Low",       "Regions with poor banking access"]],
    col_widths=[1.4, 1.8, 1.1, 0.8, 1.37])

spacer(doc,3)
comparison_cards(doc,[
    ("SWIFT\nStrengths",
     ["Trusted by 11,000+ institutions","Works for any currency pair","Highly regulated & audited",
      "Supports all financial message types","Deep compliance integration","50 years of reliability"],
     LIGHT_GREY, SWIFT_BLUE),
    ("SWIFT\nWeaknesses",
     ["1–5 day settlement time","Expensive — multiple fee layers","Complex correspondent chains",
      "Limited payment visibility (pre-GPI)","High compliance overhead","Legacy MT format"],
     RGBColor(0xFF,0xEB,0xEB), SWIFT_RED),
    ("Future\nOutlook",
     ["ISO 20022 migration underway","SWIFT GPI improving speed","Central bank CBDC pilots active",
      "API-based connectivity (Messaging API)","Real-time payments growing","SWIFT adapting to compete"],
     RGBColor(0xE8,0xF8,0xEE), SWIFT_GREEN),
])

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 10 — ISO 20022: THE FUTURE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
divider(doc)
spacer(doc,4)
section_banner(doc,"10","ISO 20022 — The Next Generation of Financial Messaging",
    "Why the entire industry is migrating from MT messages to XML — and what it means for you", SWIFT_GREEN)

body_p(doc,
    "The SWIFT MT message format has served the industry well since 1977 — but it is a relic "
    "of a teletype era. Fields are limited in length, data is unstructured, and machines "
    "struggle to parse free-text narratives. ISO 20022 (pronounced 'ISO twenty-oh-twenty-two') "
    "is the new global standard built for the digital age.")
spacer(doc,4)

tbl = doc.add_table(rows=1,cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
lc = tbl.cell(0,0); rc = tbl.cell(0,1)
lc.width = Inches(3.1); rc.width = Inches(3.37)
shade(lc, RGBColor(0xE8,0xF4,0xFF)); shade(rc, RGBColor(0xE8,0xFF,0xF0))
thin_border(lc,'0070C0'); thin_border(rc,'007A4C')

lh = lc.paragraphs[0]; lh.paragraph_format.left_indent = Pt(6)
lh.paragraph_format.space_before = Pt(6)
add_run(lh, "MT Format (Legacy FIN)\n", SWIFT_LIGHT, 9.5, bold=True)
add_run(lh, "Plain text, fixed field codes, character limits", DARK_GREY, 8, italic=True)
p = lc.add_paragraph(); p.paragraph_format.left_indent = Pt(6)
p.paragraph_format.space_before = Pt(4)
for line in [":32A:260422GBP27000,",
             ":50K:/GB12BARC...",
             ":59:/JP123456...",
             "    TOKYO SUPPLIER CO",
             ":70:/INV/2026-SUPPLY",
             ":71A:SHA"]:
    add_run(p, line+"\n", SWIFT_BLUE, 8, name='Courier New')
for con in ["✗  35-char remittance limit","✗  Unstructured free text","✗  Hard for machines to parse","✗  No legal entity identifiers","✗  Limited transaction purpose codes"]:
    lp = lc.add_paragraph(); lp.paragraph_format.left_indent = Pt(6); lp.paragraph_format.space_after = Pt(2)
    add_run(lp, con, SWIFT_RED, 8.5)
lc.add_paragraph().paragraph_format.space_after = Pt(6)

rh = rc.paragraphs[0]; rh.paragraph_format.left_indent = Pt(6)
rh.paragraph_format.space_before = Pt(6)
add_run(rh, "MX Format (ISO 20022 XML)\n", SWIFT_GREEN, 9.5, bold=True)
add_run(rh, "Structured XML with rich, unlimited data fields", DARK_GREY, 8, italic=True)
p = rc.add_paragraph(); p.paragraph_format.left_indent = Pt(6)
p.paragraph_format.space_before = Pt(4)
for line in ["<CdtTrfTxInf>",
             "  <Amt Ccy='GBP'>27000</Amt>",
             "  <Cdtr><Nm>Tokyo Supplier Co</Nm>",
             "  <Dbtr><LEI>5493...</LEI></Dbtr>",
             "  <Purp><Cd>SUPP</Cd></Purp>",
             "  <RmtInf>Invoice 2026-04-22</RmtInf>",
             "</CdtTrfTxInf>"]:
    add_run(p, line+"\n", SWIFT_GREEN, 8, name='Courier New')
for pro in ["✓  Unlimited remittance information","✓  Fully structured & machine-readable","✓  Legal Entity Identifier (LEI) fields","✓  Rich transaction purpose codes","✓  Better AML / fraud screening"]:
    rp = rc.add_paragraph(); rp.paragraph_format.left_indent = Pt(6); rp.paragraph_format.space_after = Pt(2)
    add_run(rp, pro, SWIFT_GREEN, 8.5)
rc.add_paragraph().paragraph_format.space_after = Pt(6)
spacer(doc,4)

highlight_box(doc,"ISO 20022\nMIGRATION",
    "SWIFT's cross-border ISO 20022 migration (CBPR+) began in March 2023, with a coexistence "
    "period running until November 2025 during which both MT and MX messages are supported. "
    "After 2025, MX becomes the primary standard. All financial institutions must upgrade their "
    "systems to send and receive ISO 20022 messages — this is one of the largest infrastructure "
    "changes in banking history.",
    CREAM, SWIFT_GREEN, WHITE)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 11 — KEY TAKEAWAYS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
divider(doc)
spacer(doc,4)
section_banner(doc,"11","Key Takeaways",
    "Everything you need to remember from this article in one place", SWIFT_BLUE)

body_p(doc, "If you take only ten things from this article, make it these:", color=SWIFT_BLUE)
spacer(doc,4)

takeaways = [
    (SWIFT_BLUE,  "01", "SWIFT does NOT move money",
     "It moves secure messages — instructions that tell banks to debit and credit accounts."),
    (SWIFT_TEAL,  "02", "Every bank has a BIC code",
     "The BIC (SWIFT code) is the bank's unique address on the global network — like an email address for financial institutions."),
    (SWIFT_GREEN, "03", "Correspondent banks are the connectors",
     "When two banks have no direct relationship, correspondent banks act as intermediaries — each adding fees and time."),
    (SWIFT_GOLD,  "04", "Nostro/Vostro accounts fund the system",
     "Cross-border payments work because banks pre-fund accounts at each other. No money physically crosses borders."),
    (SWIFT_RED,   "05", "MT103 is the most important message type",
     "It carries the customer credit transfer instruction — the foundation of every international wire payment."),
    (PURPLE,      "06", "MT940 / CAMT.053 reconcile the books",
     "End-of-day account statement messages allow banks to reconcile Nostro accounts and confirm every transaction."),
    (ORANGE,      "07", "Speed and cost depend on routing",
     "Fewer hops = faster, cheaper payments. Major currency corridors are always faster than exotic pairs."),
    (SWIFT_BLUE,  "08", "Security is multi-layered",
     "PKI certificates, mTLS, digital signatures, sanctions screening, and the CSP programme protect every message."),
    (SWIFT_TEAL,  "09", "SWIFT GPI changed transparency forever",
     "The UETR tracking reference means you can now track cross-border payments in real time — like a parcel."),
    (SWIFT_GREEN, "10", "ISO 20022 is the future",
     "The entire industry is migrating from legacy MT messages to rich XML-based MX messages for better data quality."),
]
n = len(takeaways)
# Render as 2-column grid
for i in range(0, n, 2):
    pair = takeaways[i:i+2]
    tbl  = doc.add_table(rows=1, cols=len(pair)); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j,(col,num,title,desc) in enumerate(pair):
        c = tbl.cell(0,j); shade(c, LIGHT_GREY); thin_border(c, rgb_hex(col)); c.width = Inches(6.47/2)
        p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.left_indent = Pt(8)
        add_run(p, f"{num}  ", col, 13, bold=True, name='Courier New')
        add_run(p, title+"\n", col, 10, bold=True)
        p2 = c.add_paragraph(); p2.paragraph_format.left_indent = Pt(8)
        p2.paragraph_format.space_after = Pt(6)
        add_run(p2, desc, DARK_GREY, 9)
    spacer(doc,3)

# ── CLOSING CALL TO ACTION ────────────────────────────────────────────────────
spacer(doc,6)
banner(doc,
    "Found this useful? Follow for more in the Financial Systems Knowledge Series",
    "Next Article: ISO 20022 Deep Dive — Understanding CAMT, PACS & PAIN Message Families",
    SLATE, WHITE, AMBER, 11, 9)

spacer(doc,4)
tbl = doc.add_table(rows=1,cols=3); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i,(tag,bg) in enumerate([
    ("#SWIFT  #CrossBorderPayments  #FinancialSystems", SWIFT_BLUE),
    ("#ISO20022  #Remittance  #Banking  #FinTech",      SWIFT_TEAL),
    ("#Payments  #GlobalFinance  #KnowledgeSeries",     SWIFT_GREEN),
]):
    c = tbl.cell(0,i); shade(c,bg); no_border(c); c.width = Inches(6.47/3)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
    add_run(p, tag, WHITE, 7.5, italic=True)

# ── FOOTER ────────────────────────────────────────────────────────────────────
spacer(doc,4)
ft = doc.add_table(rows=1,cols=1); fc = ft.cell(0,0)
shade(fc, SWIFT_BLUE); no_border(fc)
fp = fc.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(6); fp.paragraph_format.space_after = Pt(6)
add_run(fp,
    "© 2026 Financial Systems Knowledge Series  |  SWIFT Security Officer  |  "
    "For educational purposes only — not financial or legal advice",
    MID_GREY, 7.5)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = r"c:\Subash\Learning\claude\project\Research\SWIFT\SWIFT-INTRO.docx"
doc.save(out)
print(f"Saved: {out}")
