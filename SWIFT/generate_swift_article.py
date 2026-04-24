"""
Generates SWIFT-INTRO.docx — concise 4-minute LinkedIn knowledge article.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── palette ───────────────────────────────────────────────────────────────────
B   = RGBColor(0x00,0x33,0x72)   # SWIFT Blue
LB  = RGBColor(0x00,0x70,0xC0)   # Light Blue
G   = RGBColor(0x00,0x7A,0x4C)   # Green
T   = RGBColor(0x00,0x8B,0x8B)   # Teal
GD  = RGBColor(0xF0,0xA0,0x00)   # Gold
OR  = RGBColor(0xD4,0x6A,0x00)   # Orange
RD  = RGBColor(0xC0,0x00,0x00)   # Red
PU  = RGBColor(0x5B,0x2C,0x8D)   # Purple
SL  = RGBColor(0x2C,0x3E,0x50)   # Slate
W   = RGBColor(0xFF,0xFF,0xFF)
LG  = RGBColor(0xF4,0xF6,0xF8)
MG  = RGBColor(0xBF,0xBF,0xBF)
DG  = RGBColor(0x22,0x22,0x22)
CRM = RGBColor(0xFF,0xFD,0xF2)

def _hex(c): return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

# ── xml helpers ───────────────────────────────────────────────────────────────
def _shd(cell, c):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    for x in p.findall(qn('w:shd')): p.remove(x)
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto')
    s.set(qn('w:fill'),_hex(c)); p.append(s)

def _no_bdr(cell):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        t = OxmlElement(f'w:{e}')
        t.set(qn('w:val'),'none'); t.set(qn('w:sz'),'0')
        t.set(qn('w:space'),'0'); t.set(qn('w:color'),'auto')
        b.append(t)
    p.append(b)

def _bdr(cell, col='C0C0C0', sz='4'):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        t = OxmlElement(f'w:{e}')
        t.set(qn('w:val'),'single'); t.set(qn('w:sz'),sz)
        t.set(qn('w:space'),'0'); t.set(qn('w:color'),col)
        b.append(t)
    p.append(b)

def _left_bdr(cell, col, sz='12'):
    """Left-only border for pull-quote style."""
    tc = cell._tc; p = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for e in ('top','right','bottom'):
        t = OxmlElement(f'w:{e}')
        t.set(qn('w:val'),'none'); t.set(qn('w:sz'),'0')
        t.set(qn('w:space'),'0'); t.set(qn('w:color'),'auto')
        b.append(t)
    l = OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),sz)
    l.set(qn('w:space'),'0'); l.set(qn('w:color'),col)
    b.append(l); p.append(b)

# ── paragraph builders ────────────────────────────────────────────────────────
def sp(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after  = Pt(pts)

def run(para, text, color=DG, size=10.5, bold=False, italic=False, name='Calibri'):
    r = para.add_run(text)
    r.font.color.rgb=color; r.font.size=Pt(size)
    r.font.bold=bold; r.font.italic=italic; r.font.name=name
    return r

def body(doc, text, color=DG, size=10.5, italic=False, sb=3, sa=5, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent: p.paragraph_format.left_indent = Pt(indent)
    run(p, text, color, size, italic=italic)
    return p

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(5)
    run(p, text, B, 14, bold=True)
    # Blue bottom border
    pp = p._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'6')
    bt.set(qn('w:space'),'2'); bt.set(qn('w:color'),_hex(B))
    bd.append(bt); pp.append(bd)
    return p

def h2(doc, text, color=B):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run(p, text, color, 11.5, bold=True)
    return p

def bullet(doc, icon, bold_text, rest_text, icon_color=B):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent       = Pt(14)
    p.paragraph_format.first_line_indent = Pt(-14)
    run(p, icon+"  ", icon_color, 9.5, bold=True)
    if bold_text: run(p, bold_text+"  ", B, 10, bold=True)
    run(p, rest_text, DG, 10)

def divider(doc, color=_hex(MG)):
    p = doc.add_paragraph()
    pp = p._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'6')
    bt.set(qn('w:space'),'1'); bt.set(qn('w:color'),color)
    bd.append(bt); pp.append(bd)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)

# ── component builders ────────────────────────────────────────────────────────
def banner(doc, title, sub=None, bg=B, fg=W, sub_fg=None):
    t = doc.add_table(rows=1,cols=1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0,0); _shd(c,bg); _no_bdr(c)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4 if sub else 14)
    run(p, title, fg, 20, bold=True)
    if sub:
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(10)
        run(p2, sub, sub_fg or RGBColor(0xE9,0xC4,0x6A), 10, italic=True)
    sp(doc, 6)

def section_tag(doc, number, title, color=B):
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    nc = t.cell(0,0); tc = t.cell(0,1)
    nc.width = Inches(0.45); tc.width = Inches(6.02)
    _shd(nc, GD); _no_bdr(nc)
    _shd(tc, color); _no_bdr(tc)
    np = nc.paragraphs[0]; np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np.paragraph_format.space_before = Pt(7); np.paragraph_format.space_after = Pt(7)
    run(np, number, B, 11.5, bold=True)
    tp = tc.paragraphs[0]; tp.paragraph_format.left_indent = Pt(10)
    tp.paragraph_format.space_before = Pt(7); tp.paragraph_format.space_after = Pt(7)
    run(tp, title, W, 11.5, bold=True)
    sp(doc,4)

def callout(doc, label, text, bg=CRM, label_bg=GD):
    t = doc.add_table(rows=1,cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = t.cell(0,0); rc = t.cell(0,1)
    lc.width = Inches(0.82); rc.width = Inches(5.65)
    _shd(lc,label_bg); _no_bdr(lc); _shd(rc,bg); _no_bdr(rc)
    lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(8); lp.paragraph_format.space_after = Pt(8)
    run(lp, label, W, 8.5, bold=True)
    rp = rc.paragraphs[0]; rp.paragraph_format.left_indent = Pt(8)
    rp.paragraph_format.space_before = Pt(8); rp.paragraph_format.space_after = Pt(8)
    run(rp, text, DG, 9.5)
    sp(doc,4)

def pull_quote(doc, text, color=B):
    t = doc.add_table(rows=1,cols=1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0,0); _shd(c,LG); _left_bdr(c,_hex(GD),'18')
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Pt(14)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    run(p, "\u201c "+text+" \u201d", color, 12, italic=True)
    sp(doc,4)

def diagram_ref(doc, filename, caption):
    t = doc.add_table(rows=1,cols=1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0,0); _shd(c, RGBColor(0xE8,0xF0,0xFE))
    _bdr(c, _hex(LB), '6')
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    run(p, "📊  ", B, 14)
    run(p, "See Diagram: ", LB, 9.5, bold=True)
    run(p, filename, B, 9.5, bold=True, name='Courier New')
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    run(p2, caption, RGBColor(0x44,0x55,0x66), 8.5, italic=True)
    sp(doc,4)

def step_flow(doc, steps):
    """Horizontal step tiles."""
    n = len(steps)
    t = doc.add_table(rows=1,cols=n*2-1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    W_node = (6.47 - (n-1)*0.22) / n
    for i,(label,sub,color) in enumerate(steps):
        c = t.cell(0,i*2); _shd(c,color); _no_bdr(c); c.width = Inches(W_node)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(3)
        run(p, label, W, 8.5, bold=True)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(7)
        run(p2, sub, RGBColor(0xE9,0xC4,0x6A), 7.5, italic=True)
        if i < n-1:
            a = t.cell(0,i*2+1); _shd(a,W); _no_bdr(a); a.width = Inches(0.22)
            ap = a.paragraphs[0]; ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ap.paragraph_format.space_before = Pt(7)
            run(ap,"▶",GD,12,bold=True)
    sp(doc,4)

def msg_table(doc, rows_data):
    t = doc.add_table(rows=1+len(rows_data), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.style = 'Table Grid'
    for i,(h,bg) in enumerate(zip(["Message","What It Does","When Used"],
                                  [B,B,B])):
        c = t.rows[0].cells[i]; _shd(c,B); _bdr(c,_hex(B))
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        run(p, h, W, 9, bold=True)
    for ri,(code,desc,when,color) in enumerate(rows_data):
        bg = W if ri%2==0 else LG
        for ci,(val,co) in enumerate([(code,color),(desc,DG),(when,RGBColor(0x44,0x55,0x66))]):
            c = t.rows[ri+1].cells[ci]; _shd(c,bg); _bdr(c)
            p = c.paragraphs[0]
            p.paragraph_format.left_indent  = Pt(4)
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            run(p, val, co, 9, bold=(ci==0))
    for row in t.rows:
        row.cells[0].width = Inches(1.1)
        row.cells[1].width = Inches(3.3)
        row.cells[2].width = Inches(2.07)
    sp(doc,4)

def takeaway_grid(doc, items):
    pairs = [items[i:i+2] for i in range(0,len(items),2)]
    for pair in pairs:
        n = len(pair)
        t = doc.add_table(rows=1,cols=n); t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for j,(num,title,desc,color) in enumerate(pair):
            c = t.cell(0,j); _shd(c,LG); _bdr(c,_hex(color),'8'); c.width = Inches(6.47/2)
            p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.left_indent = Pt(8)
            run(p, num+"  ", color, 13, bold=True, name='Courier New')
            run(p, title+"\n", color, 9.5, bold=True)
            p2 = c.add_paragraph(); p2.paragraph_format.left_indent = Pt(8)
            p2.paragraph_format.space_after = Pt(7)
            run(p2, desc, DG, 9)
        sp(doc,3)

# ─────────────────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.27); sec.page_height   = Inches(11.69)
sec.left_margin   = Inches(0.9);  sec.right_margin  = Inches(0.9)
sec.top_margin    = Inches(0.7);  sec.bottom_margin = Inches(0.7)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

# ── COVER ─────────────────────────────────────────────────────────────────────
banner(doc,
    "Cross-Border Remittance & SWIFT",
    "Knowledge Series  ·  4-minute read  ·  Beginner → Intermediate")

# Meta row
mt = doc.add_table(rows=1,cols=4); mt.alignment = WD_TABLE_ALIGNMENT.LEFT
for i,(l,v,bg) in enumerate([
    ("Series",    "Financial Systems 101", SL),
    ("Level",     "Beginner → Intermediate", G),
    ("Read Time", "≈ 4 minutes", LB),
    ("Topic",     "Cross-Border Payments", OR),
]):
    c = mt.cell(0,i); _shd(c,bg); _no_bdr(c); c.width = Inches(6.47/4)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(2)
    run(p, l+"\n", RGBColor(0xE9,0xC4,0x6A), 7, bold=True)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(5)
    run(p2, v, W, 8.5)
sp(doc,8)

# ── HOOK ─────────────────────────────────────────────────────────────────────
body(doc,
    "You send £27,000 from London to Tokyo. Within 1–3 days the money appears "
    "in your supplier's account on the other side of the world. "
    "But how? No van drove it there. No plane carried it. "
    "What actually happened?",
    B, 11.5, italic=True, sb=4, sa=4)

body(doc,
    "The answer involves a 50-year-old network, a chain of international banks, "
    "and secure messages flying across a private internet you've never heard of. "
    "This article explains it — simply, visually, and completely.",
    DG, 10.5, sb=2, sa=8)

divider(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1 — THE PROBLEM
# ═════════════════════════════════════════════════════════════════════════════
section_tag(doc,"01","The Problem: Banks Don't Talk to Each Other")

body(doc,
    "Your bank in London has no direct relationship with a Japanese bank in Tokyo. "
    "They speak different systems, hold different currencies, and operate under different regulators. "
    "Before 1977, banks sent international payment instructions by telex — "
    "free-form text messages that were slow, error-prone, and easily forged.")

sp(doc,3)

t = doc.add_table(rows=1,cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
for i,(hdr,pts,bg,fg) in enumerate([
    ("Telex Era Problems",
     ["No standard message format","Messages easily forged","No delivery confirmation",
      "Manual re-keying at every bank","Processing took days to weeks"],
     RGBColor(0xFF,0xEB,0xEB), RD),
    ("What the World Needed",
     ["A universal message standard","Secure, authenticated messaging","Automated error-free routing",
      "Instant delivery confirmation","A network all banks could join"],
     RGBColor(0xE8,0xF8,0xEE), G),
]):
    c = t.cell(0,i); _shd(c,bg); _bdr(c,_hex(fg),'8'); c.width = Inches(6.47/2)
    p = c.paragraphs[0]; p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(4)
    run(p, hdr+"\n", fg, 9.5, bold=True)
    for pt in pts:
        pp = c.add_paragraph(); pp.paragraph_format.left_indent = Pt(8)
        pp.paragraph_format.space_after = Pt(2)
        run(pp, ("✗  " if i==0 else "✓  ")+pt, fg if i==1 else DG, 9)
    c.add_paragraph().paragraph_format.space_after = Pt(6)
sp(doc,4)

body(doc,
    "In 1973, 239 banks from 15 countries met in Brussels with one mission: "
    "build a better way. The result launched four years later — SWIFT.",
    SL, 10, italic=True)

divider(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2 — WHAT IS SWIFT
# ═════════════════════════════════════════════════════════════════════════════
section_tag(doc,"02","What is SWIFT?")

pull_quote(doc,
    "SWIFT does not move money. It moves messages — secure instructions that tell "
    "banks to move money.")

body(doc,
    "SWIFT (Society for Worldwide Interbank Financial Telecommunication) is a "
    "member-owned cooperative headquartered in Belgium. It operates a private, "
    "encrypted network — SWIFTNet — connecting over 11,000 financial institutions "
    "in 200+ countries. Every day, 44 million messages flow across it, representing "
    "over $5 trillion in transactions.")

sp(doc,4)

# Stats strip
st = doc.add_table(rows=1,cols=4); st.alignment = WD_TABLE_ALIGNMENT.LEFT
for i,(val,lbl,bg) in enumerate([
    ("11,000+","Banks connected",         B),
    ("44 M",   "Messages / day",          LB),
    ("$5 T+",  "Daily transaction value", G),
    ("200+",   "Countries covered",       T),
]):
    c = st.cell(0,i); _shd(c,bg); _no_bdr(c); c.width = Inches(6.47/4)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    run(p, val, W, 17, bold=True)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    run(p2, lbl, RGBColor(0xE9,0xC4,0x6A), 7.5, italic=True)
sp(doc,6)

h2(doc,"The BIC Code — A Bank's Global Address")
body(doc,
    "Every SWIFT member has a Bank Identifier Code (BIC) — an 8 or 11-character "
    "code that uniquely identifies it on the network. Think of it as the bank's "
    "email address. You cannot send a payment without it.")

sp(doc,3)
diagram_ref(doc,
    "SWIFT-01-BIC-Code.drawio",
    "BIC code anatomy — four components, colour-coded")

divider(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 3 — HOW A PAYMENT TRAVELS
# ═════════════════════════════════════════════════════════════════════════════
section_tag(doc,"03","How a Cross-Border Payment Travels")

body(doc,
    "Follow £27,000 from London to Tokyo. Every hop is a SWIFT message — "
    "no money physically crosses any border.")
sp(doc,4)

step_flow(doc,[
    ("YOU\nLondon",      "Payment instruction",    SL),
    ("BARCLAYS\nBARCGB22","Ordering Bank",         B),
    ("JP MORGAN\nCHASUS33","Correspondent Bank",  LB),
    ("MIZUHO\nMHCBJPJT", "Beneficiary Bank",      G),
    ("SUPPLIER\nTokyo",  "Funds credited",         OR),
])

body(doc,"What happens at each step:", DG, 10)
sp(doc,2)

for step, bold, detail in [
    ("①", "You instruct your bank.",
     "You initiate the transfer in your banking app — amount, beneficiary BIC, account number."),
    ("②", "Barclays sends a SWIFT MT103.",
     "An MT103 (Customer Credit Transfer) message flies to JP Morgan — Barclays' USD correspondent — "
     "via the SWIFTNet private network."),
    ("③", "JP Morgan routes onward.",
     "JP Morgan debits Barclays' pre-funded account (Nostro account) and sends another MT103 to Mizuho in Tokyo."),
    ("④", "Mizuho credits the supplier.",
     "Mizuho receives the message, converts to JPY, and deposits funds into the supplier's account."),
    ("⑤", "Statements reconcile the books.",
     "At end-of-day, Mizuho sends an MT940 account statement back to Barclays confirming every movement."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent  = Pt(10)
    run(p, step+"  ", GD, 11, bold=True)
    run(p, bold+"  ", B, 10, bold=True)
    run(p, detail, DG, 9.5)

sp(doc,4)
callout(doc, "KEY INSIGHT",
    "No money physically crosses borders. What moves are SWIFT messages — instructions "
    "to debit and credit Nostro accounts that banks already hold at each other. "
    "The £27,000 never leaves the banking system; it is simply re-assigned.", LG, B)

diagram_ref(doc,
    "SWIFT-02-Payment-Flow.drawio",
    "End-to-end payment journey with message labels and Nostro account illustration")

divider(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 4 — THE NETWORK
# ═════════════════════════════════════════════════════════════════════════════
section_tag(doc,"04","The SWIFT Network — How It Stays Secure")

body(doc,
    "SWIFTNet is a private IP network — entirely separate from the public internet. "
    "Every message is digitally signed with PKI certificates (proving who sent it) "
    "and encrypted with TLS (so only the recipient can read it). "
    "SWIFT validates each message before delivery.")
sp(doc,3)

# Security flow horizontal
step_flow(doc,[
    ("Compose\nMessage",   "Bank operator",   SL),
    ("Sign with\nPKI Key", "Authentication",  B),
    ("Encrypt\nTLS",       "Confidentiality", LB),
    ("SWIFT\nValidates",   "SWIFTNet private",G),
    ("Deliver +\nACK/NAK", "Confirmed",       T),
])

diagram_ref(doc,
    "SWIFT-03-Network-Architecture.drawio",
    "SWIFT hub-and-spoke architecture — banks, services (FIN / FileAct / InterAct), and the network layer")

divider(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 5 — MESSAGE TYPES
# ═════════════════════════════════════════════════════════════════════════════
section_tag(doc,"05","The Four Messages You Need to Know")

body(doc,
    "SWIFT has over 200 message types across 9 categories. "
    "For cross-border payments and treasury operations, these four are essential:", sb=2, sa=6)

msg_table(doc,[
    ("MT103",    "Customer Credit Transfer — the international wire payment instruction",
     "Every consumer/corporate cross-border payment",              B),
    ("MT202",    "Bank-to-Bank Credit Transfer — settlement between correspondent banks",
     "Interbank settlement leg of correspondent payments",         LB),
    ("MT940",    "Account Statement (legacy) — end-of-day Nostro reconciliation",
     "Daily Nostro account reconciliation",                        T),
    ("CAMT.053", "Account Statement (ISO 20022 XML) — richer, structured data format",
     "Modern ISO 20022 statement — replacing MT940",               G),
])

divider(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 6 — WHY DELAYS HAPPEN
# ═════════════════════════════════════════════════════════════════════════════
section_tag(doc,"06","Why Cross-Border Payments Take 1–5 Days",T)

body(doc,
    "Each bank in the correspondent chain runs its own compliance checks, "
    "currency conversions, and cut-off windows. More hops = more time and fees.")
sp(doc,3)

for icon,bold_text,detail in [
    ("🔍","Sanctions screening","Every bank runs OFAC, UN, and EU checks — flagged payments go for manual review."),
    ("⏰","Cut-off times","Miss a bank's daily cut-off (often 14:00–15:00 local) and the payment waits until tomorrow."),
    ("💱","Currency conversion","Each FX conversion adds spread cost and potential delay."),
    ("🌍","Time zones","New York and Tokyo overlap for just 1–2 hours per day."),
    ("🏦","More hops","Each correspondent hop adds fees ($10–$30) and hours."),
]:
    bullet(doc, icon, bold_text, detail, GD)

sp(doc,3)
callout(doc,"SWIFT GPI",
    "Launched 2017. Every GPI payment has a unique tracking reference (UETR) — like a parcel "
    "tracking number. Over 70% of SWIFT cross-border payments now arrive within 30 minutes. "
    "Transparency was the missing piece — GPI added it.", CRM, T)

divider(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 7 — THE FUTURE
# ═════════════════════════════════════════════════════════════════════════════
section_tag(doc,"07","The Future: ISO 20022",G)

body(doc,
    "The entire industry is migrating from legacy MT text messages (SWIFT FIN) "
    "to ISO 20022 XML messages — richer structured data, better fraud screening, "
    "machine-readable remittance information, and multi-language support.")
sp(doc,3)

t = doc.add_table(rows=1,cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
for i,(hdr,sample,pts,bg,fc) in enumerate([
    ("MT940 (Legacy FIN — plain text)",
     ":60F:C260422GBP1250000,\n:61:2604220422D15000,NCHK/REF123\n:86:VENDOR PAYMENT",
     ["35-character field limits","Free text — hard to parse","No structured remittance data"],
     RGBColor(0xE8,0xF0,0xFF), LB),
    ("CAMT.053 (ISO 20022 — XML)",
     "<Ntry><Amt Ccy='GBP'>15000</Amt>\n<CdtDbtInd>DBIT</CdtDbtInd>\n<RmtInf>Invoice INV-2026-042</RmtInf>",
     ["Unlimited structured data","Machine-readable XML","Full remittance & LEI fields"],
     RGBColor(0xE8,0xF8,0xF0), G),
]):
    c = t.cell(0,i); _shd(c,bg); _bdr(c,_hex(fc),'8'); c.width = Inches(6.47/2)
    p = c.paragraphs[0]; p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(4)
    run(p, hdr+"\n", fc, 9, bold=True)
    p2 = c.add_paragraph(); p2.paragraph_format.left_indent = Pt(8)
    p2.paragraph_format.space_after = Pt(4)
    run(p2, sample, RGBColor(0x33,0x55,0x88) if i==0 else RGBColor(0x00,0x55,0x33), 8, name='Courier New')
    for pt in pts:
        pp = c.add_paragraph(); pp.paragraph_format.left_indent = Pt(8); pp.paragraph_format.space_after = Pt(2)
        run(pp, ("✗  " if i==0 else "✓  ")+pt, fc, 8.5)
    c.add_paragraph().paragraph_format.space_after = Pt(6)
sp(doc,4)

body(doc,
    "SWIFT's CBPR+ (Cross-Border Payments Reporting+) migration ran from 2023 "
    "with a coexistence period until November 2025. After that, MX is primary. "
    "Every bank must upgrade — this is the largest infrastructure change in banking history.",
    SL, 9.5, italic=True)

divider(doc)

# ═════════════════════════════════════════════════════════════════════════════
# TAKEAWAYS
# ═════════════════════════════════════════════════════════════════════════════
section_tag(doc,"  ✓","5 Things to Remember", G)

sp(doc,3)
takeaway_grid(doc,[
    ("01","SWIFT moves messages, not money",
     "It is a secure messaging network connecting 11,000+ banks — not a bank itself.", B),
    ("02","BIC codes are bank addresses",
     "Every institution has a unique 8/11-character BIC for routing on the SWIFT network.", LB),
    ("03","Correspondent banks are the connectors",
     "Banks with no direct relationship route payments through pre-arranged intermediaries.", T),
    ("04","Nostro accounts fund the system",
     "No money crosses borders. Pre-funded inter-bank accounts are debited and credited.", G),
    ("05","ISO 20022 is the future",
     "The industry is replacing legacy MT messages with rich XML — better data, better compliance.", OR),
])

sp(doc,6)

# ── CLOSING CTA ───────────────────────────────────────────────────────────────
cta = doc.add_table(rows=1,cols=1); cta.alignment = WD_TABLE_ALIGNMENT.LEFT
cc = cta.cell(0,0); _shd(cc,SL); _no_bdr(cc)
p1 = cc.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before = Pt(12); p1.paragraph_format.space_after = Pt(4)
run(p1,"Found this useful? Share it with your network.", W, 11, bold=True)
p2 = cc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
run(p2,"Next in series: SWIFT Message Types Deep Dive — MT103, MT202, MT940 & CAMT.053 Explained",
    RGBColor(0xE9,0xC4,0x6A), 9, italic=True)
p3 = cc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(12)
run(p3,
    "#SWIFT  #CrossBorderPayments  #FinancialSystems  #ISO20022  #Banking  #KnowledgeSeries",
    MG, 8, italic=True)
sp(doc,4)

# ── FOOTER ────────────────────────────────────────────────────────────────────
ft = doc.add_table(rows=1,cols=1); fc2 = ft.cell(0,0)
_shd(fc2,B); _no_bdr(fc2)
fp = fc2.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(6); fp.paragraph_format.space_after = Pt(6)
run(fp,
    "Financial Systems Knowledge Series  ·  v1.0  ·  April 2026  "
    "·  For educational purposes only — not financial advice",
    MG, 7.5)

out = r"c:\Subash\Learning\claude\project\Research\SWIFT\SWIFT-INTRO.docx"
doc.save(out)
print(f"Saved Word doc: {out}")
